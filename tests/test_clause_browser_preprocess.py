from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from app.clause_browser.backend.preprocess import build_clause_browser_corpora_by_release


def build_sample_doc(path: Path, title: str = "System architecture for the 5G System (5GS); Stage 2 (Release 18)") -> None:
    doc = Document()
    cover = doc.add_table(rows=3, cols=1)
    cover.cell(0, 0).text = "3GPP TS 23.501 V18.12.0 (2025-12)"
    cover.cell(1, 0).text = "Technical Specification"
    cover.cell(2, 0).text = title
    doc.add_paragraph("1 Scope", style="Heading 1")
    doc.add_paragraph("This clause introduces the scope.")
    doc.save(path)


def test_build_clause_browser_corpora_by_release_writes_grouped_outputs(tmp_path: Path) -> None:
    first = tmp_path / "Specs" / "2024-09" / "Rel-18" / "23501-a.docx"
    second = tmp_path / "Specs" / "2024-12" / "Rel-18" / "29512-a.docx"
    first.parent.mkdir(parents=True, exist_ok=True)
    second.parent.mkdir(parents=True, exist_ok=True)
    build_sample_doc(first)
    build_sample_doc(second, title="5G System; Technical Realization of Service Based Architecture; Stage 3")

    output_root = tmp_path / "artifacts" / "clause_browser_corpora"
    media_dir = tmp_path / "artifacts" / "clause_browser_media"
    summary = build_clause_browser_corpora_by_release(
        inputs=[str(tmp_path / "Specs")],
        output_root=output_root,
        media_dir=media_dir,
        workers=2,
    )

    assert "2024-09/Rel-18" in summary
    assert "2024-12/Rel-18" in summary

    first_output = output_root / "2024-09" / "Rel-18" / "clause_browser_corpus.jsonl"
    second_output = output_root / "2024-12" / "Rel-18" / "clause_browser_corpus.jsonl"
    assert first_output.exists()
    assert second_output.exists()

    first_records = [json.loads(line) for line in first_output.read_text(encoding="utf-8").splitlines() if line.strip()]
    second_records = [json.loads(line) for line in second_output.read_text(encoding="utf-8").splitlines() if line.strip()]
    assert {record["release_data"] for record in first_records} == {"2024-09"}
    assert {record["release_data"] for record in second_records} == {"2024-12"}
