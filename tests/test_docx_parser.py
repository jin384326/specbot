from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from parser.docx_clause_parser import (
    DocxClauseParser,
    SpecMetadata,
    clean_table_matrix,
    dedupe_duplicate_cell_texts_preserve_order,
    dedupe_duplicate_brackets,
    dedupe_repeated_lines_and_semicolon_lists,
    linearized_row_pairs,
    normalize_table_cell_text,
    remove_redundant_brackets_matching_outside,
    split_relative_clause_heading,
    split_clause_heading,
    table_to_linearized_text,
)


def build_sample_doc(path: Path) -> None:
    doc = Document()
    cover = doc.add_table(rows=3, cols=1)
    cover.cell(0, 0).text = "3GPP TS 23.501 V18.12.0 (2025-12)"
    cover.cell(1, 0).text = "Technical Specification"
    cover.cell(2, 0).text = (
        "3rd Generation Partnership Project; "
        "Technical Specification Group Services and System Aspects; "
        "System architecture for the 5G System (5GS); "
        "Stage 2 (Release 18)"
    )
    doc.add_paragraph("1 Scope", style="Heading 1")
    doc.add_paragraph("This clause introduces the scope.")
    doc.add_paragraph("This clause also references 3GPP TS 23.502 and clause 4.2.1.")
    doc.add_paragraph("2 Architecture", style="Heading 1")
    for idx in range(6):
        doc.add_paragraph(f"Paragraph {idx} describing session management and access control.")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Parameter"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "SSC mode"
    table.cell(1, 1).text = "1"
    table.cell(2, 0).text = "Always-on PDU"
    table.cell(2, 1).text = "Supported"
    doc.save(path)


def test_clause_heading_parsing(tmp_path: Path) -> None:
    source = tmp_path / "2025-12" / "Rel-18" / "23501-test.docx"
    source.parent.mkdir(parents=True)
    build_sample_doc(source)
    parser = DocxClauseParser()
    records = parser.parse(source, SpecMetadata(spec_no="23501", spec_title="System architecture"))

    clause_docs = [record for record in records if record.doc_type == "clause_doc"]
    assert [record.clause_id for record in clause_docs] == ["1", "2"]
    assert clause_docs[0].clause_title == "Scope"
    assert clause_docs[1].clause_title == "Architecture"


def test_table_matrix_and_linearized_text(tmp_path: Path) -> None:
    source = tmp_path / "table.docx"
    build_sample_doc(source)
    doc = Document(source)
    matrix = clean_table_matrix(doc.tables[1])
    assert matrix[0] == ["Parameter", "Value"]
    linearized = table_to_linearized_text(matrix, "Architecture")
    assert "row 1: Parameter: SSC mode; Value: 1" in linearized
    assert "Table Architecture" in linearized


def test_dedupe_duplicate_brackets_in_cell() -> None:
    assert dedupe_duplicate_brackets("a [PDU Session request] b [PDU Session request]") == "a [PDU Session request] b"
    assert dedupe_duplicate_brackets("[] []") == "[]"


def test_remove_redundant_brackets_when_repeated_outside() -> None:
    assert (
        normalize_table_cell_text("PDU Session Establishment Request [PDU Session Establishment Request]")
        == "PDU Session Establishment Request"
    )


def test_linearized_row_skips_consecutive_duplicate_merged_cells() -> None:
    header = ["Message", "Message", "Direction"]
    row = ["Foo request", "Foo request", "UE to CN"]
    pairs = linearized_row_pairs(header, row)
    assert pairs == [("Message", "Foo request"), ("Direction", "UE to CN")]


def test_dedupe_semicolon_and_line_repetition_in_list_like_cell() -> None:
    cell = (
        "PDU Session Establishment Request; PDU Session Establishment Request; "
        "PDU Session Modification Request"
    )
    assert (
        dedupe_repeated_lines_and_semicolon_lists(cell)
        == "PDU Session Establishment Request; PDU Session Modification Request"
    )
    multiline = "The SMF shall send X.\nThe SMF shall send X.\nThe AMF shall receive Y."
    out = dedupe_repeated_lines_and_semicolon_lists(multiline, min_line_len=15)
    assert out.count("The SMF shall send X.") == 1
    assert "The AMF shall receive Y." in out


def test_normalize_table_cell_applies_list_dedupe() -> None:
    assert (
        normalize_table_cell_text("Note A; Note A; Note B")
        == "Note A; Note B"
    )


def test_clean_table_matrix_collapses_horizontal_and_vertical_merges(tmp_path: Path) -> None:
    path = tmp_path / "merged.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=4)
    t.cell(0, 0).text = "H0"
    t.cell(0, 1).text = "H1"
    t.cell(0, 2).text = "H2"
    t.cell(0, 3).text = "H3"
    t.cell(1, 0).text = "A"
    t.cell(1, 1).text = "B"
    t.cell(1, 2).text = "C"
    t.cell(1, 3).text = "D"
    t.cell(0, 0).merge(t.cell(0, 1))
    t.cell(1, 0).merge(t.cell(1, 2))
    doc.save(path)

    doc2 = Document(path)
    matrix = clean_table_matrix(doc2.tables[0])
    assert matrix[0][:3] == ["H0 H1", "H2", "H3"]
    assert matrix[1][:2] == ["A B C", "D"]

    path_v = tmp_path / "merged_v.docx"
    doc_v = Document()
    tv = doc_v.add_table(rows=3, cols=2)
    tv.cell(0, 0).text = "Name"
    tv.cell(0, 1).text = "Val"
    tv.cell(1, 0).text = "Span"
    tv.cell(1, 1).text = "x"
    tv.cell(2, 0).text = "y"
    tv.cell(2, 1).text = "z"
    tv.cell(1, 0).merge(tv.cell(2, 0))
    doc_v.save(path_v)
    matrix_v = clean_table_matrix(Document(path_v).tables[0])
    assert matrix_v[1] == ["Span y", "x"]
    assert matrix_v[2] == ["", "z"]
    linearized = table_to_linearized_text(matrix_v, "T")
    assert linearized.count("Span") == 1


def test_clean_table_matrix_prefers_majority_row_width_over_inflated_column_count() -> None:
    class FakeCell:
        def __init__(self, text: str) -> None:
            self.text = text

    class FakeRow:
        def __init__(self, values: list[str]) -> None:
            self.cells = [FakeCell(value) for value in values]

    class FakeTable:
        def __init__(self) -> None:
            self._column_count = 4
            self.rows = [
                FakeRow(["Data type", "Reference", "Comments"]),
                FakeRow(["DurationSec", "3GPP TS 29.571 [7]", "Time value in seconds"]),
                FakeRow(["BitRate", "3GPP TS 29.571 [7]", "", ""]),
            ]

    matrix = clean_table_matrix(FakeTable())

    assert matrix == [
        ["Data type", "Reference", "Comments"],
        ["DurationSec", "3GPP TS 29.571 [7]", "Time value in seconds"],
        ["BitRate", "3GPP TS 29.571 [7]", ""],
    ]


def test_dedupe_duplicate_cell_texts_collapses_merged_and_interleaved_cells() -> None:
    label = "Odd/even indication (octet 4) Bit"
    assert dedupe_duplicate_cell_texts_preserve_order([label] * 10) == [label]
    assert dedupe_duplicate_cell_texts_preserve_order(["A", "A", "B", "B", "B"]) == ["A", "B"]
    assert dedupe_duplicate_cell_texts_preserve_order([label, "", label, "", label]) == [label]
    assert dedupe_duplicate_cell_texts_preserve_order(["A", "B", "A"]) == ["A", "B"]


def test_passage_splitting(tmp_path: Path) -> None:
    source = tmp_path / "passages.docx"
    build_sample_doc(source)
    parser = DocxClauseParser(passage_char_limit=90, passage_paragraph_limit=2, min_passage_chars=10)
    records = parser.parse(source, SpecMetadata(spec_no="23501"))

    passages = [record for record in records if record.doc_type == "passage_doc" and record.clause_id == "2"]
    assert len(passages) >= 2
    assert passages[0].paragraph_start_index <= passages[0].paragraph_end_index


def test_annex_heading_and_false_positive_guard() -> None:
    assert split_clause_heading("Annex A (informative) Example flows") == ("Annex A", "Example flows")
    assert split_clause_heading("5.28a.3 Topology Information for TSN TN") == ("5.28a.3", "Topology Information for TSN TN")
    assert split_clause_heading("D.1 Topology information for transport") == ("D.1", "Topology information for transport")
    assert split_clause_heading("5G architecture") is None


def test_parser_keeps_mixed_alphanumeric_clause_ids(tmp_path: Path) -> None:
    source = tmp_path / "mixed-clause.docx"
    doc = Document()
    cover = doc.add_table(rows=3, cols=1)
    cover.cell(0, 0).text = "3GPP TS 23.501 V18.12.0 (2025-12)"
    cover.cell(1, 0).text = "Technical Specification"
    cover.cell(2, 0).text = "System architecture for the 5G System (5GS); Stage 2 (Release 18)"
    doc.add_paragraph("5.28a.3 Topology Information for TSN TN", style="Heading 3")
    doc.add_paragraph("Topology details.")
    doc.save(source)

    parser = DocxClauseParser()
    records = parser.parse(source, SpecMetadata(spec_no="23501"))

    clause_docs = [record for record in records if record.doc_type == "clause_doc"]
    assert [record.clause_id for record in clause_docs] == ["5.28a.3"]
    assert clause_docs[0].clause_title == "Topology Information for TSN TN"


def test_parser_treats_h6_style_as_heading_but_not_b1(tmp_path: Path) -> None:
    source = tmp_path / "h6-clause.docx"
    doc = Document()
    doc.styles.add_style("B1", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("H6", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("4.11.0a Impacts to EPS Procedures", style="Heading 3")
    doc.add_paragraph("4.11.0a.2 Interaction with PCC", style="Heading 4")
    doc.add_paragraph("4.11.0a.2.1 Nested H6 clause", style="H6")
    doc.add_paragraph("Nested body.")
    doc.add_paragraph("8a to 8c. This is body text only.", style="B1")
    doc.save(source)

    parser = DocxClauseParser()
    records = parser.parse(source, SpecMetadata(spec_no="23502"))

    clause_ids = [record.clause_id for record in records if record.doc_type == "clause_doc"]
    assert "4.11.0a.2.1" in clause_ids
    assert "8a" not in clause_ids


def test_parser_restores_relative_heading_numbers_from_parent_prefix(tmp_path: Path) -> None:
    source = tmp_path / "relative-heading.docx"
    doc = Document()
    doc.add_paragraph("7.2.14 Modify Bearer Command and Failure Indication", style="Heading 3")
    doc.add_paragraph(".1 Modify Bearer Command", style="Heading 4")
    doc.add_paragraph("Body for .1")
    doc.add_paragraph(".2 Modify Bearer Failure Indication", style="Heading 4")
    doc.add_paragraph("Body for .2")
    doc.save(source)

    parser = DocxClauseParser()
    records = parser.parse(source, SpecMetadata(spec_no="29274"))

    clause_docs = [record for record in records if record.doc_type == "clause_doc"]
    clause_ids = [record.clause_id for record in clause_docs]
    assert clause_ids == ["7.2.14.1", "7.2.14.2"]
    assert clause_docs[0].clause_title == "Modify Bearer Command"
    assert clause_docs[-1].clause_title == "Modify Bearer Failure Indication"


def test_split_relative_clause_heading() -> None:
    assert split_relative_clause_heading(".2 Modify Bearer Failure Indication") == ("2", "Modify Bearer Failure Indication")
    assert split_relative_clause_heading(".0a.5 Impacts to EPS Procedures") == ("0a.5", "Impacts to EPS Procedures")
    assert split_relative_clause_heading("7.2.14.2 Modify Bearer Failure Indication") is None


def test_parser_restores_missing_numeric_heading_from_previous_sibling(tmp_path: Path) -> None:
    source = tmp_path / "implicit-heading.docx"
    doc = Document()
    doc.add_paragraph("7.4 CS Fallback and SRVCC related messages", style="Heading 2")
    doc.add_paragraph("7.4.1 Suspend Notification", style="Heading 3")
    doc.add_paragraph("Body 1")
    doc.add_paragraph("7.4.2 Suspend Acknowledge", style="Heading 3")
    doc.add_paragraph("Body 2")
    doc.add_paragraph("7.4.3 Resume Notification", style="Heading 3")
    doc.add_paragraph("Body 3")
    doc.add_paragraph("Resume Acknowledge", style="Heading 3")
    doc.add_paragraph("Body 4")
    doc.save(source)

    parser = DocxClauseParser()
    records = parser.parse(source, SpecMetadata(spec_no="29274"))

    clause_docs = [record for record in records if record.doc_type == "clause_doc"]
    clause_ids = [record.clause_id for record in clause_docs]
    assert "7.4.4" in clause_ids
    restored = next(record for record in clause_docs if record.clause_id == "7.4.4")
    assert restored.clause_title == "Resume Acknowledge"


def test_parser_prefixes_immediate_child_title_when_parent_has_no_body(tmp_path: Path) -> None:
    source = tmp_path / "heading-only.docx"
    doc = Document()
    doc.add_paragraph("5.37.8 UE power saving management", style="Heading 3")
    doc.add_paragraph("5.37.8.1 General", style="Heading 4")
    doc.add_paragraph("General body.")
    doc.save(source)

    parser = DocxClauseParser(prefix_direct_child_title_from_empty_parent=True)
    records = parser.parse(source, SpecMetadata(spec_no="23501"))

    clause_docs = [record for record in records if record.doc_type == "clause_doc"]
    child = next(record for record in clause_docs if record.clause_id == "5.37.8.1")

    assert all(record.clause_id != "5.37.8" for record in clause_docs)
    assert child.clause_title == "UE power saving management : General"
    assert child.text == "General body."


def test_parser_does_not_prefix_grandchild_title_from_heading_only_grandparent(tmp_path: Path) -> None:
    source = tmp_path / "heading-only-grandchild.docx"
    doc = Document()
    doc.add_paragraph("5.37.8 UE power saving management", style="Heading 3")
    doc.add_paragraph("5.37.8.1 General", style="Heading 4")
    doc.add_paragraph("General body.")
    doc.add_paragraph("5.37.8.1.1 Details", style="Heading 5")
    doc.add_paragraph("Details body.")
    doc.save(source)

    parser = DocxClauseParser(prefix_direct_child_title_from_empty_parent=True)
    records = parser.parse(source, SpecMetadata(spec_no="23501"))

    clause_docs = [record for record in records if record.doc_type == "clause_doc"]
    child = next(record for record in clause_docs if record.clause_id == "5.37.8.1")
    grandchild = next(record for record in clause_docs if record.clause_id == "5.37.8.1.1")

    assert child.clause_title == "UE power saving management : General"
    assert grandchild.clause_title == "Details"


def test_parser_default_keeps_browser_titles_without_parent_prefix(tmp_path: Path) -> None:
    source = tmp_path / "browser-title.docx"
    doc = Document()
    doc.add_paragraph("5.37.8 UE power saving management", style="Heading 3")
    doc.add_paragraph("5.37.8.1 General", style="Heading 4")
    doc.add_paragraph("General body.")
    doc.save(source)

    parser = DocxClauseParser()
    records = parser.parse(source, SpecMetadata(spec_no="23501"))

    clause_docs = [record for record in records if record.doc_type == "clause_doc"]
    child = next(record for record in clause_docs if record.clause_id == "5.37.8.1")

    assert all(record.clause_id != "5.37.8" for record in clause_docs)
    assert child.clause_title == "General"


def test_parser_prefix_uses_immediate_parent_title_without_accumulating_ancestor_prefix(tmp_path: Path) -> None:
    source = tmp_path / "nested-empty-parents.docx"
    doc = Document()
    doc.add_paragraph("5.37 High level features", style="Heading 2")
    doc.add_paragraph("5.37.8 UE power saving management", style="Heading 3")
    doc.add_paragraph("5.37.8.1 General", style="Heading 4")
    doc.add_paragraph("General body.")
    doc.save(source)

    parser = DocxClauseParser(prefix_direct_child_title_from_empty_parent=True)
    records = parser.parse(source, SpecMetadata(spec_no="23501"))

    clause_docs = [record for record in records if record.doc_type == "clause_doc"]
    child = next(record for record in clause_docs if record.clause_id == "5.37.8.1")

    assert child.clause_title == "UE power saving management : General"




def test_spec_title_and_release_data_are_inferred_from_docx_and_path(tmp_path: Path) -> None:
    source = tmp_path / "2025-12" / "Rel-18" / "23501-demo.docx"
    source.parent.mkdir(parents=True)
    build_sample_doc(source)
    parser = DocxClauseParser()
    records = parser.parse(source, SpecMetadata(spec_no="23501"))

    first_clause = next(record for record in records if record.doc_type == "clause_doc")
    assert first_clause.spec_title == "System architecture for the 5G System (5GS); Stage 2 (Release 18)"
    assert first_clause.release_data == "2025-12"
    assert first_clause.stage_hint == "Stage 2"


def test_unknown_stage_is_marked_as_else(tmp_path: Path) -> None:
    source = tmp_path / "unknown-stage.docx"
    doc = Document()
    cover = doc.add_table(rows=3, cols=1)
    cover.cell(0, 0).text = "3GPP TS 29.999 V1.0.0 (2025-12)"
    cover.cell(1, 0).text = "Technical Specification"
    cover.cell(2, 0).text = "Some specification without explicit stage"
    doc.add_paragraph("1 Scope", style="Heading 1")
    doc.add_paragraph("Body text")
    doc.save(source)

    parser = DocxClauseParser()
    records = parser.parse(source, SpecMetadata(spec_no="29999"))

    first_clause = next(record for record in records if record.doc_type == "clause_doc")
    assert first_clause.stage_hint == "else"


def test_parser_skips_annex_and_change_history_sections(tmp_path: Path) -> None:
    source = tmp_path / "skip-annex.docx"
    doc = Document()
    cover = doc.add_table(rows=3, cols=1)
    cover.cell(0, 0).text = "3GPP TS 23.501 V18.12.0 (2025-12)"
    cover.cell(1, 0).text = "Technical Specification"
    cover.cell(2, 0).text = "System architecture for the 5G System (5GS); Stage 2 (Release 18)"
    doc.add_paragraph("1 Scope", style="Heading 1")
    doc.add_paragraph("Normal body text.")
    doc.add_paragraph("Annex A (informative) Example flows", style="Heading 1")
    doc.add_paragraph("This annex should not become a corpus record.")
    doc.add_paragraph("Annex B (informative) Change history", style="Heading 1")
    doc.add_paragraph("This change history should not become a corpus record.")
    doc.save(source)

    parser = DocxClauseParser()
    records = parser.parse(source, SpecMetadata(spec_no="23501"))

    clause_titles = [record.clause_title for record in records if record.doc_type == "clause_doc"]
    assert clause_titles == ["Scope"]


def test_normal_annex_sentence_is_not_misparsed_as_annex_heading(tmp_path: Path) -> None:
    source = tmp_path / "annex-false-positive.docx"
    doc = Document()
    cover = doc.add_table(rows=3, cols=1)
    cover.cell(0, 0).text = "3GPP TS 23.501 V18.12.0 (2025-12)"
    cover.cell(1, 0).text = "Technical Specification"
    cover.cell(2, 0).text = "System architecture for the 5G System (5GS); Stage 2 (Release 18)"
    doc.add_paragraph("5.27.2 Parent clause", style="Heading 3")
    doc.add_paragraph("Body.")
    doc.add_paragraph("5.27.2.2 Previous clause", style="Heading 4")
    doc.add_paragraph("Body.")
    doc.add_paragraph("Annex I describe how the traffic pattern information is determined.", style="Normal")
    doc.add_paragraph("5.27.2.3 TSC Assistance Container determination by TSCTSF", style="Heading 4")
    doc.add_paragraph("Body.")
    doc.save(source)

    parser = DocxClauseParser()
    records = parser.parse(source, SpecMetadata(spec_no="23501"))

    clause_docs = [record for record in records if record.doc_type == "clause_doc"]
    clause_ids = [record.clause_id for record in clause_docs]
    assert "Annex I" not in clause_ids
    target = next(record for record in clause_docs if record.clause_id == "5.27.2.3")
    assert target.parent_clause_id == "5.27.2"
    assert target.clause_path == ["5.27.2", "5.27.2.3"]


def test_normal_numbered_sentence_is_not_misparsed_as_clause_heading(tmp_path: Path) -> None:
    source = tmp_path / "numbered-false-positive.docx"
    doc = Document()
    cover = doc.add_table(rows=3, cols=1)
    cover.cell(0, 0).text = "3GPP TS 23.501 V18.12.0 (2025-12)"
    cover.cell(1, 0).text = "Technical Specification"
    cover.cell(2, 0).text = "System architecture for the 5G System (5GS); Stage 2 (Release 18)"
    doc.add_paragraph("5.15.18 Parent clause", style="Heading 3")
    doc.add_paragraph("Body.")
    doc.add_paragraph("5.15.18.2 Previous clause", style="Heading 4")
    doc.add_paragraph("Body.")
    doc.add_paragraph(
        "3 If the UE has overlapping areas between non-allowed area, a cell inside the NS-AoS, then the non-allowed area restriction applies.",
        style="Normal",
    )
    doc.add_paragraph(
        "5.15.18.3 Network based monitoring and enforcement of Network Slice Area of Service not matching deployed Tracking Areas",
        style="Heading 4",
    )
    doc.add_paragraph("Body.")
    doc.save(source)

    parser = DocxClauseParser()
    records = parser.parse(source, SpecMetadata(spec_no="23501"))

    clause_docs = [record for record in records if record.doc_type == "clause_doc"]
    clause_ids = [record.clause_id for record in clause_docs]
    assert "3" not in clause_ids
    target = next(record for record in clause_docs if record.clause_id == "5.15.18.3")
    assert target.parent_clause_id == "5.15.18"
    assert target.clause_path == ["5.15.18", "5.15.18.3"]


def test_numeric_clause_after_annex_does_not_inherit_annex_in_clause_path(tmp_path: Path) -> None:
    source = tmp_path / "annex-reset.docx"
    doc = Document()
    cover = doc.add_table(rows=3, cols=1)
    cover.cell(0, 0).text = "3GPP TS 23.501 V18.12.0 (2025-12)"
    cover.cell(1, 0).text = "Technical Specification"
    cover.cell(2, 0).text = "System architecture for the 5G System (5GS); Stage 2 (Release 18)"
    doc.add_paragraph("Annex I (informative) Extra material", style="Heading 1")
    doc.add_paragraph("Annex content.")
    doc.add_paragraph("5.37 Main clause after annex", style="Heading 2")
    doc.add_paragraph("Main body.")
    doc.add_paragraph("5.37.8 Child clause", style="Heading 3")
    doc.add_paragraph("Child body.")
    doc.add_paragraph("5.37.8.3 Leaf clause", style="Heading 4")
    doc.add_paragraph("Leaf body.")
    doc.save(source)

    parser = DocxClauseParser()
    records = parser.parse(source, SpecMetadata(spec_no="23501"))

    clause_docs = [record for record in records if record.doc_type == "clause_doc"]
    leaf_clause = next(record for record in clause_docs if record.clause_id == "5.37.8.3")
    assert leaf_clause.clause_path == ["5.37", "5.37.8", "5.37.8.3"]
    assert leaf_clause.parent_clause_id == "5.37.8"


def test_annex_letter_clause_is_parsed_as_clause(tmp_path: Path) -> None:
    source = tmp_path / "annex-letter-clause.docx"
    doc = Document()
    cover = doc.add_table(rows=3, cols=1)
    cover.cell(0, 0).text = "3GPP TS 23.501 V18.12.0 (2025-12)"
    cover.cell(1, 0).text = "Technical Specification"
    cover.cell(2, 0).text = "System architecture for the 5G System (5GS); Stage 2 (Release 18)"
    doc.add_paragraph("Annex D (informative) Transport aspects", style="Heading 1")
    doc.add_paragraph("Annex body.")
    doc.add_paragraph("D.1 Topology information for transport", style="Heading 2")
    doc.add_paragraph("Clause body.")
    doc.add_paragraph("D.1.1 Transport topology details", style="Heading 3")
    doc.add_paragraph("Nested clause body.")
    doc.save(source)

    parser = DocxClauseParser()
    records = parser.parse(source, SpecMetadata(spec_no="23501"))

    clause_docs = [record for record in records if record.doc_type == "clause_doc"]
    d1_clause = next(record for record in clause_docs if record.clause_id == "D.1")
    d11_clause = next(record for record in clause_docs if record.clause_id == "D.1.1")
    assert d1_clause.clause_path == ["Annex D", "D.1"]
    assert d1_clause.parent_clause_id == "Annex D"
    assert d11_clause.clause_path == ["Annex D", "D.1", "D.1.1"]
    assert d11_clause.parent_clause_id == "D.1"
