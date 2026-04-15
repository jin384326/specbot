from __future__ import annotations

import asyncio
import json
import base64
import zipfile
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from fastapi import HTTPException
from fastapi.routing import APIRoute

from app.clause_browser.api import (
    ClauseBrowserConfig,
    ExportRequest,
    LLMActionRequest,
    SpecbotQueryRequest,
    create_router,
)
from app.clause_browser.backend.render_parser import RichDocxClauseParser
from app.clause_browser.backend.board_api import BoardCreatePayload, BoardLockPayload, BoardUpdatePayload
from app.clause_browser.preprocess import build_clause_browser_corpus
from app.clause_browser.repository import ClauseRepository
from app.clause_browser.server import ClauseBrowserSettings, create_app
from app.clause_browser.services import (
    DocxExportService,
    LLMActionQueueFullError,
    LLMActionService,
    MarkdownExportService,
    SpecbotQueryDefaults,
    SpecbotQueryService,
    sanitize_file_stem,
)
from app.specbot_query_server import PersistentSpecbotQueryEngine


def write_corpus(path: Path) -> None:
    records = [
        {
            "doc_type": "clause_doc",
            "spec_no": "23501",
            "spec_title": "System architecture",
            "release": "Rel-18",
            "release_data": "2025-12",
            "clause_id": "5",
            "clause_title": "Session management",
            "parent_clause_id": "",
            "clause_path": ["5"],
            "text": "Session management overview.",
            "source_file": str(path.parent / "23501.docx"),
            "order_in_source": 1,
        },
        {
            "doc_type": "clause_doc",
            "spec_no": "23501",
            "spec_title": "System architecture",
            "release": "Rel-18",
            "release_data": "2025-12",
            "clause_id": "5.1",
            "clause_title": "PDU session procedures",
            "parent_clause_id": "5",
            "clause_path": ["5", "5.1"],
            "text": "Detailed clause body.",
            "source_file": str(path.parent / "23501.docx"),
            "order_in_source": 2,
        },
        {
            "doc_type": "clause_doc",
            "spec_no": "23501",
            "spec_title": "System architecture",
            "release": "Rel-18",
            "release_data": "2025-12",
            "clause_id": "5.1.1",
            "clause_title": "Request handling",
            "parent_clause_id": "5.1",
            "clause_path": ["5", "5.1", "5.1.1"],
            "text": "Request handling details.",
            "source_file": str(path.parent / "23501.docx"),
            "order_in_source": 3,
        },
        {
            "doc_type": "clause_doc",
            "spec_no": "24501",
            "spec_title": "Core network procedures",
            "release": "Rel-17",
            "release_data": "2024-12",
            "clause_id": "1",
            "clause_title": "Scope",
            "parent_clause_id": "",
            "clause_path": ["1"],
            "text": "Another document.",
            "source_file": str(path.parent / "24501.docx"),
            "order_in_source": 1,
        },
        {
            "doc_type": "clause_doc",
            "spec_no": "23502",
            "spec_title": "Procedures for the 5G System (5GS); Stage 2 (Release 18)",
            "release": "Rel-18",
            "release_data": "2025-12",
            "clause_id": "4.2.2.2",
            "clause_title": "PDU session establishment",
            "parent_clause_id": "",
            "clause_path": ["4.2.2.2"],
            "text": "",
            "source_file": str(path.parent / "23502.docx"),
            "order_in_source": 1,
            "blocks": [
                {
                    "type": "table",
                    "rows": [
                        ["Field", "Description"],
                        ["redundantPduSessionInfo", "Contains RSN and PDU Session Pair ID"],
                    ],
                }
            ],
        },
    ]
    with path.open("w", encoding="utf-8") as handle:
        for record in records:
            handle.write(json.dumps(record) + "\n")


def write_docx_files(tmp_path: Path) -> None:
    image_path = tmp_path / "tiny.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+a7VQAAAAASUVORK5CYII="
        )
    )

    doc = Document()
    doc.add_paragraph("5 Session management", style="Heading 1")
    doc.add_paragraph("Session management overview.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Mode"
    table.cell(1, 1).text = "SSC1"
    doc.add_paragraph().add_run().add_picture(str(image_path))
    caption_style = doc.styles.add_style("TF", WD_STYLE_TYPE.PARAGRAPH)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 5-1: Example figure caption", style="TF")
    doc.add_paragraph("5.1 PDU session procedures", style="Heading 2")
    doc.add_paragraph("Detailed clause body.")
    doc.add_paragraph("5.1.1 Request handling", style="Heading 3")
    doc.add_paragraph("Request handling details.")
    doc.save(tmp_path / "23501.docx")

    doc2 = Document()
    doc2.add_paragraph("1 Scope", style="Heading 1")
    doc2.add_paragraph("Another document.")
    doc2.save(tmp_path / "24501.docx")


def write_docx_with_body_numbered_paragraph(tmp_path: Path) -> Path:
    doc = Document()
    doc.styles.add_style("B1", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("5 Session management", style="Heading 1")
    doc.add_paragraph("Overview text.")
    doc.add_paragraph("3 Packet Delay Budget (including Core Network Packet Delay Budget);", style="B1")
    doc.add_paragraph("This line must stay in clause 5.")
    path = tmp_path / "numbered-body.docx"
    doc.save(path)
    return path


def write_docx_with_outline_level_heading(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("8.2.22 Offending IE", style="Heading 2")
    doc.add_paragraph("Offending IE body.")
    paragraph = doc.add_paragraph("8.2.222 N6 Jitter Measurement")
    p_pr = paragraph._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "2")
    p_pr.append(outline)
    doc.add_paragraph("The N6 Jitter Measurement IE contains a N6 jitter measurement.")
    path = tmp_path / "outline-level-heading.docx"
    doc.save(path)
    return path


def write_docx_with_merged_table(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("5 Session management", style="Heading 1")
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Group"
    table.cell(0, 1).text = "Field"
    table.cell(0, 2).text = "Value"
    table.cell(1, 0).text = "Session"
    table.cell(1, 1).text = "Mode"
    table.cell(1, 2).text = "SSC1"
    table.cell(2, 0).text = "Will be merged away"
    table.cell(2, 1).text = "Timer"
    table.cell(2, 2).text = "30s"
    table.cell(1, 0).merge(table.cell(2, 0))
    path = tmp_path / "23501.docx"
    doc.save(path)
    return path


def write_docx_with_indented_paragraph(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("5 Session management", style="Heading 1")
    paragraph = doc.add_paragraph(
        "1a. (UE initiated modification) The UE initiates the PDU Session Modification procedure."
    )
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)
    path = tmp_path / "23501-indented-paragraph.docx"
    doc.save(path)
    return path


def write_docx_with_sized_image(tmp_path: Path) -> Path:
    image_path = tmp_path / "sized.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+a7VQAAAAASUVORK5CYII="
        )
    )
    doc = Document()
    doc.add_paragraph("5 Session management", style="Heading 1")
    doc.add_paragraph("Sized image example.")
    doc.add_paragraph().add_run().add_picture(str(image_path), width=Inches(4))
    path = tmp_path / "sized-image.docx"
    doc.save(path)
    return path


def build_app(tmp_path: Path):
    write_docx_files(tmp_path)
    export_dir = tmp_path / "exports"
    browser_corpus_path = tmp_path / "browser-corpus.jsonl"
    build_clause_browser_corpus(
        inputs=[str(tmp_path / "23501.docx"), str(tmp_path / "24501.docx")],
        output_path=browser_corpus_path,
        media_dir=tmp_path / "media",
    )
    with browser_corpus_path.open("a", encoding="utf-8") as handle:
        handle.write(
            json.dumps(
                {
                    "doc_type": "clause_doc",
                    "content_kind": "clause",
                    "doc_id": "23502:clause:4.2.2.2",
                    "spec_no": "23502",
                    "spec_title": "Procedures for the 5G System (5GS); Stage 2 (Release 18)",
                    "clause_id": "4.2.2.2",
                    "clause_title": "PDU session establishment",
                    "parent_clause_id": "",
                    "clause_path": ["4.2.2.2"],
                    "text": "",
                    "source_file": str(tmp_path / "23502.docx"),
                    "order_in_source": 1,
                    "blocks": [
                        {
                            "type": "table",
                            "rows": [
                                ["Field", "Description"],
                                ["redundantPduSessionInfo", "Contains RSN and PDU Session Pair ID"],
                            ],
                        }
                    ],
                }
            )
            + "\n"
        )
    return create_app(
        ClauseBrowserSettings(
            project_root=tmp_path,
            corpus_path=browser_corpus_path,
            export_dir=export_dir,
            media_dir=tmp_path / "media",
            cors_origins=(),
            llm_provider="mock",
            llm_model="mock-model",
            llm_max_concurrent_requests=2,
            llm_max_queued_requests=5,
            languages=(("ko", "Korean"), ("en", "English")),
            specbot_query_api_url="http://127.0.0.1:8010",
        ),
        llm_service=LLMActionService(provider="mock", model="mock-model"),
    )


def get_endpoint(app, path: str):
    for route in app.routes:
        if isinstance(route, APIRoute) and route.path == path:
            return route.endpoint
    raise AssertionError(f"Route not found: {path}")


def get_endpoint_with_method(app, path: str, method: str):
    for route in app.routes:
        if isinstance(route, APIRoute) and route.path == path and method.upper() in route.methods:
            return route.endpoint
    raise AssertionError(f"Route not found: {method} {path}")


def test_documents_endpoint_lists_specs(tmp_path: Path) -> None:
    app = build_app(tmp_path)
    endpoint = get_endpoint(app, "/api/clause-browser/documents")

    payload = endpoint()
    assert payload["success"] is True
    assert [item["specNo"] for item in payload["data"]["items"]] == ["23501", "23502", "24501"]


def test_translation_split_preserves_paragraph_boundaries() -> None:
    first = "A" * 7000
    second = "B" * 6990
    text = f"{first}\n\n{second}"

    chunks = LLMActionService._split_translation_text(text, limit=12000)

    assert len(chunks) == 2
    assert chunks[0] == first
    assert chunks[1] == second


def test_documents_endpoint_can_filter_by_clause_and_body_text_before_document_selection(tmp_path: Path) -> None:
    app = build_app(tmp_path)
    endpoint = get_endpoint(app, "/api/clause-browser/documents")

    by_clause_title = endpoint("", "request handling")["data"]["items"]
    assert [item["specNo"] for item in by_clause_title] == ["23501"]

    by_body_text = endpoint("", "another document")["data"]["items"]
    assert [item["specNo"] for item in by_body_text] == ["24501"]

    by_table_text = endpoint("", "RSN")["data"]["items"]
    assert [item["specNo"] for item in by_table_text] == ["23502"]


def test_clause_repository_sorts_clauses_by_clause_path_not_source_order(tmp_path: Path) -> None:
    corpus_path = tmp_path / "browser-corpus.jsonl"
    records = [
        {
            "doc_type": "clause_doc",
            "spec_no": "23502",
            "spec_title": "Procedures for the 5G System (5GS); Stage 2",
            "release": "Rel-18",
            "release_data": "2025-12",
            "clause_id": "4.2.2",
            "clause_title": "Registration Management procedures",
            "parent_clause_id": "",
            "clause_path": ["4", "4.2", "4.2.2"],
            "text": "",
            "source_file": str(tmp_path / "23502.docx"),
            "order_in_source": 115,
        },
        {
            "doc_type": "clause_doc",
            "spec_no": "23502",
            "spec_title": "Procedures for the 5G System (5GS); Stage 2",
            "release": "Rel-18",
            "release_data": "2025-12",
            "clause_id": "4.2.2.2",
            "clause_title": "Registration procedures",
            "parent_clause_id": "4.2.2",
            "clause_path": ["4", "4.2", "4.2.2", "4.2.2.2"],
            "text": "",
            "source_file": str(tmp_path / "23502.docx"),
            "order_in_source": 118,
        },
        {
            "doc_type": "clause_doc",
            "spec_no": "23502",
            "spec_title": "Procedures for the 5G System (5GS); Stage 2",
            "release": "Rel-18",
            "release_data": "2025-12",
            "clause_id": "4.2.2.1",
            "clause_title": "General",
            "parent_clause_id": "4.2.2",
            "clause_path": ["4", "4.2", "4.2.2", "4.2.2.1"],
            "text": "",
            "source_file": str(tmp_path / "23502.docx"),
            "order_in_source": 127,
        },
        {
            "doc_type": "clause_doc",
            "spec_no": "23502",
            "spec_title": "Procedures for the 5G System (5GS); Stage 2",
            "release": "Rel-18",
            "release_data": "2025-12",
            "clause_id": "4.2.2.2.1",
            "clause_title": "General",
            "parent_clause_id": "4.2.2.2",
            "clause_path": ["4", "4.2", "4.2.2", "4.2.2.2", "4.2.2.2.1"],
            "text": "",
            "source_file": str(tmp_path / "23502.docx"),
            "order_in_source": 119,
        },
    ]
    corpus_path.write_text("".join(json.dumps(item) + "\n" for item in records), encoding="utf-8")
    repository = ClauseRepository(corpus_path, load_workers=1)

    payload = repository.list_clauses("23502", include_all=True, release_data="2025-12", release="Rel-18")
    clause_ids = [item.clause_id for item in payload if item.clause_id.startswith("4.2.2")]

    assert clause_ids == ["4.2.2", "4.2.2.1", "4.2.2.2", "4.2.2.2.1"]


def test_clause_repository_subtree_children_follow_clause_path_order(tmp_path: Path) -> None:
    corpus_path = tmp_path / "browser-corpus.jsonl"
    records = [
        {
            "doc_type": "clause_doc",
            "spec_no": "23502",
            "spec_title": "Procedures for the 5G System (5GS); Stage 2",
            "release": "Rel-18",
            "release_data": "2025-12",
            "clause_id": "4.2.2",
            "clause_title": "Registration Management procedures",
            "parent_clause_id": "",
            "clause_path": ["4", "4.2", "4.2.2"],
            "text": "",
            "source_file": str(tmp_path / "23502.docx"),
            "order_in_source": 115,
        },
        {
            "doc_type": "clause_doc",
            "spec_no": "23502",
            "spec_title": "Procedures for the 5G System (5GS); Stage 2",
            "release": "Rel-18",
            "release_data": "2025-12",
            "clause_id": "4.2.2.3",
            "clause_title": "Deregistration procedures",
            "parent_clause_id": "4.2.2",
            "clause_path": ["4", "4.2", "4.2.2", "4.2.2.3"],
            "text": "",
            "source_file": str(tmp_path / "23502.docx"),
            "order_in_source": 120,
        },
        {
            "doc_type": "clause_doc",
            "spec_no": "23502",
            "spec_title": "Procedures for the 5G System (5GS); Stage 2",
            "release": "Rel-18",
            "release_data": "2025-12",
            "clause_id": "4.2.2.2",
            "clause_title": "Registration procedures",
            "parent_clause_id": "4.2.2",
            "clause_path": ["4", "4.2", "4.2.2", "4.2.2.2"],
            "text": "",
            "source_file": str(tmp_path / "23502.docx"),
            "order_in_source": 118,
        },
        {
            "doc_type": "clause_doc",
            "spec_no": "23502",
            "spec_title": "Procedures for the 5G System (5GS); Stage 2",
            "release": "Rel-18",
            "release_data": "2025-12",
            "clause_id": "4.2.2.1",
            "clause_title": "General",
            "parent_clause_id": "4.2.2",
            "clause_path": ["4", "4.2", "4.2.2", "4.2.2.1"],
            "text": "",
            "source_file": str(tmp_path / "23502.docx"),
            "order_in_source": 127,
        },
    ]
    corpus_path.write_text("".join(json.dumps(item) + "\n" for item in records), encoding="utf-8")
    repository = ClauseRepository(corpus_path, load_workers=1)

    payload = repository.get_subtree("23502", "4.2.2", release_data="2025-12", release="Rel-18").to_dict()

    assert [item["clauseId"] for item in payload["children"]] == ["4.2.2.1", "4.2.2.2", "4.2.2.3"]


def test_clause_list_search_matches_clause_body_text(tmp_path: Path) -> None:
    app = build_app(tmp_path)
    endpoint = get_endpoint(app, "/api/clause-browser/documents/{spec_no}/clauses")

    payload = endpoint("23501", "request handling details", 100)["data"]["items"]
    assert [item["clauseId"] for item in payload] == ["5.1.1"]


def test_clause_list_search_matches_table_block_text(tmp_path: Path) -> None:
    app = build_app(tmp_path)
    endpoint = get_endpoint(app, "/api/clause-browser/documents/{spec_no}/clauses")

    payload = endpoint("23502", "RSN", 100)["data"]["items"]
    assert [item["clauseId"] for item in payload] == ["4.2.2.2"]
    assert "rsn" in payload[0]["searchText"]


def test_synthesized_ancestor_clause_is_searchable(tmp_path: Path) -> None:
    export_dir = tmp_path / "exports"
    write_docx_files(tmp_path)
    browser_corpus_path = tmp_path / "browser-corpus.jsonl"
    build_clause_browser_corpus(
        inputs=[str(tmp_path / "23501.docx")],
        output_path=browser_corpus_path,
        media_dir=tmp_path / "media",
    )
    with browser_corpus_path.open("a", encoding="utf-8") as handle:
        handle.write(
            json.dumps(
                {
                    "doc_type": "clause_doc",
                    "content_kind": "clause",
                    "doc_id": "29512:clause:5.5.2.1",
                    "spec_no": "29512",
                    "spec_title": "Example spec",
                    "clause_id": "5.5.2.1",
                    "clause_title": "Description",
                    "parent_clause_id": "5.5.2",
                    "clause_path": ["5", "5.5", "5.5.2", "5.5.2.1"],
                    "text": "Leaf body",
                    "source_file": str(tmp_path / "23501.docx"),
                    "order_in_source": 10,
                    "blocks": [{"type": "paragraph", "text": "Leaf body"}],
                }
            )
            + "\n"
        )
    app = create_app(
        ClauseBrowserSettings(
            project_root=tmp_path,
            corpus_path=browser_corpus_path,
            export_dir=export_dir,
            media_dir=tmp_path / "media",
            cors_origins=(),
            llm_provider="mock",
            llm_model="mock-model",
            llm_max_concurrent_requests=2,
            llm_max_queued_requests=5,
            languages=(("ko", "Korean"), ("en", "English")),
            specbot_query_api_url="http://127.0.0.1:8010",
        )
    )
    endpoint = get_endpoint(app, "/api/clause-browser/documents/{spec_no}/clauses")
    payload = endpoint("29512", "5.5.2", 100)["data"]["items"]
    assert any(item["clauseId"] == "5.5.2" for item in payload)


def test_subtree_endpoint_returns_descendants(tmp_path: Path) -> None:
    app = build_app(tmp_path)
    endpoint = get_endpoint(app, "/api/clause-browser/documents/{spec_no}/clauses/{clause_id:path}/subtree")

    payload = endpoint("23501", "5")["data"]
    assert payload["clauseId"] == "5"
    assert [block["type"] for block in payload["blocks"]] == ["paragraph", "table", "image", "paragraph"]
    assert payload["children"][0]["clauseId"] == "5.1"
    assert payload["children"][0]["children"][0]["clauseId"] == "5.1.1"
    assert payload["blocks"][1]["rows"][1][1] == "SSC1"


def test_clause_browser_corpus_preserves_docx_image_display_size(tmp_path: Path) -> None:
    source_path = write_docx_with_sized_image(tmp_path)
    browser_corpus_path = tmp_path / "browser-corpus.jsonl"

    build_clause_browser_corpus(
        inputs=[str(source_path)],
        output_path=browser_corpus_path,
        media_dir=tmp_path / "media",
    )

    records = [json.loads(line) for line in browser_corpus_path.open(encoding="utf-8")]
    image_block = next(block for block in records[0]["blocks"] if block["type"] == "image")
    assert image_block["displayWidthPx"] == 384
    assert image_block["displayHeightPx"] == 384


def test_rich_docx_clause_parser_prefers_png_for_vector_browser_images(tmp_path: Path, monkeypatch) -> None:
    parser = RichDocxClauseParser(media_root=tmp_path / "media")
    vector_path = tmp_path / "media" / "23501" / "5" / "diagram.wmf"
    vector_path.parent.mkdir(parents=True, exist_ok=True)
    vector_path.write_bytes(b"wmf")

    def fake_convert(self, path: Path) -> Path:
        svg_path = path.with_suffix(".svg")
        svg_path.write_text('<svg xmlns="http://www.w3.org/2000/svg"></svg>', encoding="utf-8")
        return svg_path

    def fake_export(self, source_path: Path, output_path: Path) -> None:
        output_path.write_bytes(
            base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wn9l5QAAAAASUVORK5CYII=")
        )

    monkeypatch.setattr(RichDocxClauseParser, "_convert_vector_image_if_needed", fake_convert)
    monkeypatch.setattr(RichDocxClauseParser, "_run_inkscape_png_export", fake_export)

    browser_path = parser._prepare_browser_image(vector_path)

    assert browser_path.suffix == ".png"
    assert browser_path.name == "diagram.export.png"
    assert browser_path.exists()
    assert vector_path.with_suffix(".svg").exists()


def test_rich_docx_clause_parser_uses_configured_vector_export_dpi(tmp_path: Path, monkeypatch) -> None:
    captured: dict[str, object] = {}

    def fake_run(command, check, stdout, stderr, text, env):  # type: ignore[no-untyped-def]
        captured["command"] = command
        output_path = next(Path(part.split("=", 1)[1]) for part in command if part.startswith("--export-filename="))
        output_path.write_bytes(
            base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wn9l5QAAAAASUVORK5CYII=")
        )
        class Result:
            returncode = 0
            stdout = ""
            stderr = ""
        return Result()

    monkeypatch.setenv("SPECBOT_CLAUSE_BROWSER_VECTOR_EXPORT_DPI", "288")
    monkeypatch.setenv("SPECBOT_CLAUSE_BROWSER_VECTOR_EXPORT_SCALE", "2")
    monkeypatch.setattr("app.clause_browser.backend.render_parser.subprocess.run", fake_run)

    RichDocxClauseParser._run_inkscape_png_export(tmp_path / "diagram.wmf", tmp_path / "diagram.export.png")

    assert "--export-dpi=576" in captured["command"]


def test_numbered_body_paragraph_does_not_become_clause(tmp_path: Path) -> None:
    source_path = write_docx_with_body_numbered_paragraph(tmp_path)
    browser_corpus_path = tmp_path / "browser-corpus.jsonl"

    build_clause_browser_corpus(
        inputs=[str(source_path)],
        output_path=browser_corpus_path,
        media_dir=tmp_path / "media",
    )

    records = [json.loads(line) for line in browser_corpus_path.open(encoding="utf-8")]
    clause_ids = [record["clause_id"] for record in records]
    assert clause_ids == ["5"]
    assert "Packet Delay Budget" not in records[0]["clause_title"]
    assert "3 Packet Delay Budget (including Core Network Packet Delay Budget);" in records[0]["text"]


def test_missing_clause_returns_404(tmp_path: Path) -> None:
    app = build_app(tmp_path)
    endpoint = get_endpoint(app, "/api/clause-browser/documents/{spec_no}/clauses/{clause_id:path}/subtree")

    try:
        endpoint("23501", "9.9")
    except HTTPException as exc:
        assert exc.status_code == 404
    else:
        raise AssertionError("Expected HTTPException")


def test_outline_level_heading_becomes_clause(tmp_path: Path) -> None:
    source_path = write_docx_with_outline_level_heading(tmp_path)
    browser_corpus_path = tmp_path / "browser-corpus.jsonl"

    build_clause_browser_corpus(
        inputs=[str(source_path)],
        output_path=browser_corpus_path,
        media_dir=tmp_path / "media",
    )

    records = [json.loads(line) for line in browser_corpus_path.open(encoding="utf-8")]
    clause_ids = [record["clause_id"] for record in records]
    assert "8.2.22" in clause_ids
    assert "8.2.222" in clause_ids
    outline_record = next(record for record in records if record["clause_id"] == "8.2.222")
    assert outline_record["clause_title"] == "N6 Jitter Measurement"


def test_subtree_endpoint_preserves_table_rowspan_metadata(tmp_path: Path) -> None:
    source_path = write_docx_with_merged_table(tmp_path)
    browser_corpus_path = tmp_path / "browser-corpus.jsonl"
    with browser_corpus_path.open("w", encoding="utf-8") as handle:
        handle.write(
            json.dumps(
                {
                    "doc_type": "clause_doc",
                    "content_kind": "clause",
                    "doc_id": "23501:clause:5",
                    "spec_no": "23501",
                    "spec_title": "Merged table spec",
                    "clause_id": "5",
                    "clause_title": "Session management",
                    "parent_clause_id": "",
                    "clause_path": ["5"],
                    "text": "",
                    "source_file": str(source_path),
                    "order_in_source": 1,
                    "blocks": [],
                }
            )
            + "\n"
        )
    app = create_app(
        ClauseBrowserSettings(
            project_root=tmp_path,
            corpus_path=browser_corpus_path,
            export_dir=tmp_path / "exports",
            media_dir=tmp_path / "media",
            cors_origins=(),
            llm_provider="mock",
            llm_model="mock-model",
            llm_max_concurrent_requests=2,
            llm_max_queued_requests=5,
            languages=(("ko", "Korean"), ("en", "English")),
            specbot_query_api_url="http://127.0.0.1:8010",
        )
    )
    endpoint = get_endpoint(app, "/api/clause-browser/documents/{spec_no}/clauses/{clause_id:path}/subtree")
    payload = endpoint("23501", "5")["data"]
    table_block = next(block for block in payload["blocks"] if block["type"] == "table")
    assert "cells" in table_block
    assert table_block["cells"][1][0]["text"] == "Session Will be merged away"
    assert table_block["cells"][1][0]["rowspan"] == 2
    assert table_block["cells"][1][0]["colspan"] == 1


def test_subtree_endpoint_returns_paragraph_indent_metadata(tmp_path: Path) -> None:
    source_path = write_docx_with_indented_paragraph(tmp_path)
    browser_corpus_path = tmp_path / "browser-corpus.jsonl"

    build_clause_browser_corpus(
        inputs=[str(source_path)],
        output_path=browser_corpus_path,
        media_dir=tmp_path / "media",
    )

    app = create_app(
        ClauseBrowserSettings(
            project_root=tmp_path,
            corpus_path=browser_corpus_path,
            export_dir=tmp_path / "exports",
            media_dir=tmp_path / "media",
            cors_origins=(),
            llm_provider="mock",
            llm_model="mock-model",
            llm_max_concurrent_requests=2,
            llm_max_queued_requests=5,
            languages=(("ko", "Korean"), ("en", "English")),
            specbot_query_api_url="http://127.0.0.1:8010",
        )
    )
    endpoint = get_endpoint(app, "/api/clause-browser/documents/{spec_no}/clauses/{clause_id:path}/subtree")

    payload = endpoint("23501", "5")["data"]
    paragraph_block = next(block for block in payload["blocks"] if block["type"] == "paragraph")

    assert paragraph_block["format"]["leftIndentPx"] == 24
    assert paragraph_block["format"]["leftIndentPt"] == 18.0
    assert paragraph_block["format"]["textIndentPx"] == -24
    assert paragraph_block["format"]["textIndentPt"] == -18.0


def test_docx_export_endpoint_saves_file(tmp_path: Path) -> None:
    app = build_app(tmp_path)
    subtree_endpoint = get_endpoint(app, "/api/clause-browser/documents/{spec_no}/clauses/{clause_id:path}/subtree")
    export_endpoint = get_endpoint(app, "/api/clause-browser/exports/docx")
    subtree = subtree_endpoint("23501", "5")["data"]

    payload = export_endpoint(ExportRequest(title="Session Export", roots=[subtree]))["data"]
    exported_path = tmp_path / payload["relativePath"]
    assert exported_path.is_file()
    exported = Document(exported_path)
    headings = [p.text for p in exported.paragraphs if p.style.name.startswith("Heading")]
    assert "5 Session management" in headings
    assert "5.1 PDU session procedures" in headings


def test_docx_download_endpoint_returns_attachment(tmp_path: Path) -> None:
    app = build_app(tmp_path)
    subtree_endpoint = get_endpoint(app, "/api/clause-browser/documents/{spec_no}/clauses/{clause_id:path}/subtree")
    export_endpoint = get_endpoint(app, "/api/clause-browser/exports/docx/download")
    subtree = subtree_endpoint("23501", "5")["data"]

    response = export_endpoint(ExportRequest(title="Session Export", roots=[subtree]))

    assert response.media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert response.headers["content-disposition"] == (
        'attachment; filename="Session_Export.docx"; '
        "filename*=UTF-8''Session_Export.docx"
    )
    exported = Document(BytesIO(response.body))
    headings = [p.text for p in exported.paragraphs if p.style.name.startswith("Heading")]
    assert "5 Session management" in headings


def test_docx_download_endpoint_uses_safe_ascii_fallback_for_non_ascii_title(tmp_path: Path) -> None:
    app = build_app(tmp_path)
    subtree_endpoint = get_endpoint(app, "/api/clause-browser/documents/{spec_no}/clauses/{clause_id:path}/subtree")
    export_endpoint = get_endpoint(app, "/api/clause-browser/exports/docx/download")
    subtree = subtree_endpoint("23501", "5")["data"]

    response = export_endpoint(ExportRequest(title="새 게시글", roots=[subtree]))

    assert response.headers["content-disposition"] == (
        'attachment; filename="clause-export.docx"; '
        "filename*=UTF-8''%EC%83%88_%EA%B2%8C%EC%8B%9C%EA%B8%80.docx"
    )


def test_markdown_download_endpoint_returns_attachment(tmp_path: Path) -> None:
    app = build_app(tmp_path)
    subtree_endpoint = get_endpoint(app, "/api/clause-browser/documents/{spec_no}/clauses/{clause_id:path}/subtree")
    export_endpoint = get_endpoint(app, "/api/clause-browser/exports/markdown/download")
    subtree = subtree_endpoint("23501", "5")["data"]

    response = export_endpoint(ExportRequest(title="Session Export", roots=[subtree]))

    assert response.media_type == "text/markdown"
    assert response.headers["content-disposition"] == (
        'attachment; filename="Session_Export.md"; '
        "filename*=UTF-8''Session_Export.md"
    )
    text = response.body.decode("utf-8")
    assert "# Session Export" in text
    assert "## 23501 " in text
    assert "### 5 Session management" in text


def test_markdown_package_download_endpoint_returns_zip_with_assets(tmp_path: Path) -> None:
    media_dir = tmp_path / "media" / "23501" / "5"
    media_dir.mkdir(parents=True, exist_ok=True)
    image_path = media_dir / "tiny.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wn9l5QAAAAASUVORK5CYII="
        )
    )

    app = build_app(tmp_path)
    export_endpoint = get_endpoint(app, "/api/clause-browser/exports/markdown-package/download")
    roots = [
        {
            "key": "23501:5",
            "specNo": "23501",
            "specTitle": "System architecture",
            "clauseId": "5",
            "clauseTitle": "Session management",
            "text": "",
            "blocks": [
                {"type": "image", "src": "/clause-browser-media/23501/5/tiny.png", "alt": "Figure 1"},
            ],
            "children": [],
        }
    ]

    response = export_endpoint(ExportRequest(title="Session Export", roots=[roots[0]]))

    assert response.media_type == "application/zip"
    archive = zipfile.ZipFile(BytesIO(response.body))
    names = set(archive.namelist())
    assert "document.md" in names
    assert "assets/23501/5/tiny.png" in names
    text = archive.read("document.md").decode("utf-8")
    assert "![Figure 1](assets/23501/5/tiny.png)" in text


def test_llm_action_endpoint_returns_mock_translation(tmp_path: Path) -> None:
    app = build_app(tmp_path)
    endpoint = get_endpoint(app, "/api/clause-browser/llm-actions")

    payload = asyncio.run(
        endpoint(
            None,
            LLMActionRequest(
                actionType="translate",
                text="Session management",
                sourceLanguage="en",
                targetLanguage="ko",
                context="23501 / 5",
            )
        )
    )["data"]
    assert payload["provider"] == "mock"
    assert payload["outputText"] == "[en->ko] Session management"


def test_llm_action_rejects_empty_text(tmp_path: Path) -> None:
    app = build_app(tmp_path)
    endpoint = get_endpoint(app, "/api/clause-browser/llm-actions")

    try:
        asyncio.run(
            endpoint(
                None,
                LLMActionRequest(
                    actionType="translate",
                    text=" ",
                    sourceLanguage="en",
                    targetLanguage="ko",
                )
            )
        )
    except HTTPException as exc:
        assert exc.status_code == 400
        assert "Select text" in exc.detail
    else:
        raise AssertionError("Expected HTTPException")


def test_board_update_without_lock_does_not_mark_post_as_editing(tmp_path: Path) -> None:
    app = build_app(tmp_path)
    create_post = get_endpoint_with_method(app, "/api/clause-browser/board/posts", "POST")
    update_post = get_endpoint_with_method(app, "/api/clause-browser/board/posts/{post_id}", "PUT")
    get_post = get_endpoint_with_method(app, "/api/clause-browser/board/posts/{post_id}", "GET")

    created = create_post(
        BoardCreatePayload(
            editorId="editor-1",
            editorLabel="Editor 1",
            title="Session Export",
            releaseData="2025-12",
            release="Rel-18",
            workspaceState={"notes": []},
        )
    )["data"]
    post_id = created["postId"]

    # Simulate leaving edit mode so read-only autosave updates occur without a lock.
    release_lock = get_endpoint_with_method(app, "/api/clause-browser/board/posts/{post_id}/lock/release", "POST")
    release_lock(post_id, BoardLockPayload(editorId="editor-1", editorLabel="Editor 1"))

    updated = update_post(
        post_id,
        BoardUpdatePayload(
            editorId="viewer-1",
            editorLabel="Viewer 1",
            title="Session Export",
            workspaceState={"notes": [{"id": "n1"}]},
        ),
    )["data"]

    assert updated["lock"] is None
    fetched = get_post(post_id)["data"]
    assert fetched["lock"] is None
    assert fetched["workspaceState"]["notes"] == [{"id": "n1"}]


def test_llm_action_service_rejects_requests_when_queue_is_full() -> None:
    service = LLMActionService(
        provider="mock",
        model="mock-model",
        max_concurrent_requests=2,
        max_queued_requests=1,
    )
    original_run = service.run

    def slow_run(**kwargs):
        import time

        time.sleep(0.1)
        return original_run(**kwargs)

    service.run = slow_run

    def invoke(index: int):
        return service.run_limited(
            action_type="translate",
            text=f"text-{index}",
            source_language="en",
            target_language="ko",
        )

    with ThreadPoolExecutor(max_workers=4) as executor:
        futures = [executor.submit(invoke, index) for index in range(4)]
        results = []
        for future in futures:
            try:
                results.append(future.result())
            except Exception as exc:
                results.append(exc)

    queue_errors = [item for item in results if isinstance(item, LLMActionQueueFullError)]
    assert len(queue_errors) == 1


def test_llm_action_service_uses_selection_prompts_for_selection_scope(tmp_path: Path) -> None:
    clause_system = tmp_path / "system_prompt_clause_summary.txt"
    clause_user = tmp_path / "user_prompt_clause_summary.txt"
    selection_system = tmp_path / "system_prompt_translate.txt"
    selection_user = tmp_path / "user_prompt_translate.txt"
    clause_system.write_text("CLAUSE_SYSTEM", encoding="utf-8")
    clause_user.write_text("CLAUSE_USER {context_text}", encoding="utf-8")
    selection_system.write_text("SELECTION_SYSTEM", encoding="utf-8")
    selection_user.write_text("SELECTION_USER {context_text}", encoding="utf-8")

    service = LLMActionService(
        provider="mock",
        model="mock-model",
        system_prompt_path=clause_system,
        user_prompt_path=clause_user,
        selection_system_prompt_path=selection_system,
        selection_user_prompt_path=selection_user,
    )

    assert service._get_translation_prompt_pair("selection") == ("SELECTION_SYSTEM", "SELECTION_USER {context_text}")
    assert service._get_translation_prompt_pair("clause") == ("CLAUSE_SYSTEM", "CLAUSE_USER {context_text}")
    assert service._get_translation_prompt_pair(None) == ("CLAUSE_SYSTEM", "CLAUSE_USER {context_text}")


def test_llm_action_endpoint_returns_429_when_queue_is_full(tmp_path: Path) -> None:
    class BusyLLMActionService(LLMActionService):
        def run_limited(self, **kwargs):
            raise LLMActionQueueFullError("queue full")

    write_docx_files(tmp_path)
    browser_corpus_path = tmp_path / "browser-corpus.jsonl"
    build_clause_browser_corpus(
        inputs=[str(tmp_path / "23501.docx"), str(tmp_path / "24501.docx")],
        output_path=browser_corpus_path,
        media_dir=tmp_path / "media",
    )
    router = create_router(
        repository=ClauseRepository(browser_corpus_path),
        export_service=DocxExportService(export_dir=tmp_path / "exports", project_root=tmp_path),
        llm_service=BusyLLMActionService(),
        config=ClauseBrowserConfig(
            languages=[{"code": "ko", "label": "Korean"}, {"code": "en", "label": "English"}],
            actions=[{"type": "translate", "label": "Translate"}],
            release_scopes=[],
        ),
    )
    endpoint = next(
        route.endpoint
        for route in router.routes
        if isinstance(route, APIRoute) and route.path == "/api/clause-browser/llm-actions"
    )

    try:
        asyncio.run(
            endpoint(
                None,
                LLMActionRequest(
                    actionType="translate",
                    text="Session management",
                    sourceLanguage="en",
                    targetLanguage="ko",
                )
            )
        )
    except HTTPException as exc:
        assert exc.status_code == 429
        assert exc.detail == "queue full"
    else:
        raise AssertionError("Expected HTTPException")


def test_specbot_query_endpoint_returns_hits(tmp_path: Path) -> None:
    class FakeSpecbotQueryService(SpecbotQueryService):
        def __init__(self, project_root: Path) -> None:
            super().__init__(project_root=project_root, defaults=SpecbotQueryDefaults())
            self.last_exclude_specs = None
            self.last_exclude_clauses = None
            self.last_release_data = None
            self.last_release = None

        def run(
            self,
            query: str,
            settings: dict[str, object] | None = None,
            exclude_specs: list[str] | None = None,
            exclude_clauses: list[dict[str, object]] | None = None,
            release_data: str | None = None,
            release: str | None = None,
        ) -> dict[str, object]:
            self.last_exclude_specs = exclude_specs
            self.last_exclude_clauses = exclude_clauses
            self.last_release_data = release_data
            self.last_release = release
            return {
                "query": query,
                "settings": settings or self.defaults.to_dict(),
                "hits": [
                    {
                        "specNo": "23501",
                        "clauseId": "5.1",
                        "parentClauseId": "5",
                        "clausePath": ["5", "5.1"],
                        "textPreview": "Detailed clause body.",
                    }
                ],
                "rawResult": {"relevant_documents": []},
                "command": "python3 -m app.main iterative-query-vespa-http ...",
            }

    fake_service = FakeSpecbotQueryService(tmp_path)
    app = build_app(tmp_path)
    app = create_app(
        ClauseBrowserSettings(
            project_root=tmp_path,
            corpus_path=tmp_path / "browser-corpus.jsonl",
            export_dir=tmp_path / "exports",
            media_dir=tmp_path / "media",
            cors_origins=(),
            llm_provider="mock",
            llm_model="mock-model",
            llm_max_concurrent_requests=2,
            llm_max_queued_requests=5,
            languages=(("ko", "Korean"), ("en", "English")),
            specbot_query_api_url="http://127.0.0.1:8010",
        ),
        specbot_service=fake_service,
    )
    endpoint = get_endpoint(app, "/api/clause-browser/specbot/query")

    payload = asyncio.run(
        endpoint(
            None,
            SpecbotQueryRequest(
                query="End to End Redundant Paths",
                releaseData="2025-12",
                release="Rel-18",
                excludeSpecs=["23502"],
                excludeClauses=[{"specNo": "23501", "clauseId": "5.2"}],
            ),
        )
    )["data"]
    assert payload["query"] == "End to End Redundant Paths"
    assert payload["hits"][0]["specNo"] == "23501"
    assert fake_service.last_exclude_specs == ["23502"]
    assert fake_service.last_exclude_clauses == [{"specNo": "23501", "clauseId": "5.2"}]
    assert fake_service.last_release_data == "2025-12"
    assert fake_service.last_release == "Rel-18"


def test_clause_repository_filters_documents_by_release_scope(tmp_path: Path) -> None:
    corpus_path = tmp_path / "browser-corpus.jsonl"
    write_corpus(corpus_path)
    repository = ClauseRepository(corpus_path)

    filtered = repository.list_documents(release_data="2024-12", release="Rel-17")

    assert [item.spec_no for item in filtered] == ["24501"]


def test_specbot_query_exclusion_is_exact_spec_clause_pair() -> None:
    hits = [
        {"specNo": "23502", "clauseId": "4.2.2.2"},
        {"specNo": "29512", "clauseId": "4.2.2.2"},
        {"specNo": "23502", "clauseId": "4.2.2.3"},
    ]
    filtered = PersistentSpecbotQueryEngine._apply_exclusions(
        hits,
        exclude_specs=[],
        exclude_clauses=[{"specNo": "23502", "clauseId": "4.2.2.2"}],
    )
    assert filtered == [
        {"specNo": "29512", "clauseId": "4.2.2.2"},
        {"specNo": "23502", "clauseId": "4.2.2.3"},
    ]


def test_docx_export_service_adds_numeric_suffix(tmp_path: Path) -> None:
    export_dir = tmp_path / "exports"
    service = DocxExportService(export_dir=export_dir, project_root=tmp_path)
    roots = [
        {
            "clauseId": "5",
            "clauseTitle": "Session management",
            "text": "Body",
            "children": [],
        }
    ]

    first = service.export("My Export", roots)
    second = service.export("My Export", roots)

    assert first.file_name == "My_Export.docx"
    assert second.file_name == "My_Export_2.docx"


def test_markdown_export_service_serializes_tables_images_and_notes(tmp_path: Path) -> None:
    export_dir = tmp_path / "exports"
    service = MarkdownExportService(export_dir=export_dir, project_root=tmp_path)
    roots = [
        {
            "key": "23501:5",
            "specNo": "23501",
            "specTitle": "System architecture",
            "clauseId": "5",
            "clauseTitle": "Session management",
            "text": "",
            "blocks": [
                {"type": "paragraph", "text": "Paragraph body", "format": {}},
                {
                    "type": "table",
                    "rows": [["Column A", "Column B"], ["Value 1", "Value 2"]],
                },
                {"type": "image", "src": "/clause-browser-media/23501/5/tiny.png", "alt": "Figure 1"},
            ],
            "children": [],
        }
    ]
    notes = [
        {
            "id": "selection:1",
            "type": "selection",
            "clauseKey": "23501:5",
            "blockIndex": 0,
            "translation": "Selection memo translation",
            "sourceText": "Paragraph body",
        }
    ]

    file_name, payload = service.export_bytes("My Export", roots, notes=notes)
    text = payload.decode("utf-8")

    assert file_name == "My_Export.md"
    assert "# My Export" in text
    assert "## 23501 System architecture" in text
    assert "### 5 Session management" in text
    assert "Paragraph body" in text
    assert "| Column A | Column B |" in text
    assert "| --- | --- |" in text
    assert "![Figure 1](/clause-browser-media/23501/5/tiny.png)" in text
    assert "> Note: Selection memo translation" in text


def test_markdown_export_service_includes_notes_below_table_image_and_clause(tmp_path: Path) -> None:
    export_dir = tmp_path / "exports"
    service = MarkdownExportService(export_dir=export_dir, project_root=tmp_path)
    roots = [
        {
            "key": "23501:5",
            "specNo": "23501",
            "specTitle": "System architecture",
            "clauseId": "5",
            "clauseTitle": "Session management",
            "text": "",
            "blocks": [
                {
                    "type": "table",
                    "rows": [["Column A", "Column B"], ["Value 1", "Value 2"]],
                },
                {"type": "image", "src": "/clause-browser-media/23501/5/tiny.png", "alt": "Figure 1"},
            ],
            "children": [],
        }
    ]
    notes = [
        {
            "id": "23501:5:clause",
            "type": "clause",
            "clauseKey": "23501:5",
            "translation": "Clause memo translation",
            "sourceText": "Clause source",
        },
        {
            "id": "selection:table",
            "type": "selection",
            "clauseKey": "23501:5",
            "blockIndex": 0,
            "translation": "Table memo translation",
            "sourceText": "Value 1",
        },
        {
            "id": "selection:image",
            "type": "selection",
            "clauseKey": "23501:5",
            "blockIndex": 1,
            "translation": "Image memo translation",
            "sourceText": "Figure 1",
        },
    ]

    _file_name, payload = service.export_bytes("My Export", roots, notes=notes)
    text = payload.decode("utf-8")

    assert "| Column A | Column B |" in text
    assert "| Value 1 | Value 2 |" in text
    assert "| Value 1 | Value 2 |\n\n> Note: Table memo translation" in text
    assert "![Figure 1](/clause-browser-media/23501/5/tiny.png)\n\n> Note: Image memo translation" in text
    assert "> Note: Clause memo translation" in text


def test_markdown_export_service_packages_local_images_and_note_origins(tmp_path: Path) -> None:
    export_dir = tmp_path / "exports"
    media_dir = tmp_path / "artifacts" / "clause_browser_media" / "23501" / "5"
    media_dir.mkdir(parents=True, exist_ok=True)
    image_path = media_dir / "tiny.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wn9l5QAAAAASUVORK5CYII="
        )
    )
    service = MarkdownExportService(export_dir=export_dir, project_root=tmp_path)
    roots = [
        {
            "key": "23501:5",
            "specNo": "23501",
            "specTitle": "System architecture",
            "clauseId": "5",
            "clauseTitle": "Session management",
            "text": "",
            "blocks": [
                {
                    "type": "table",
                    "rows": [["Column A", "Column B"], ["Value 1", "Value 2"]],
                },
                {"type": "image", "src": "/clause-browser-media/23501/5/tiny.png", "alt": "Figure 1"},
            ],
            "children": [],
        }
    ]
    notes = [
        {
            "id": "selection:table",
            "type": "selection",
            "clauseKey": "23501:5",
            "blockIndex": 0,
            "rowIndex": 1,
            "cellId": "cell-2",
            "translation": "Table memo translation",
            "sourceText": "Value 2",
        },
    ]

    file_name, payload = service.export_package_bytes("My Export", roots, notes=notes)

    assert file_name == "My_Export.zip"
    archive = zipfile.ZipFile(BytesIO(payload))
    assert set(archive.namelist()) == {"document.md", "assets/23501/5/tiny.png"}
    text = archive.read("document.md").decode("utf-8")
    assert "> Note (rowIndex=1, cellId=cell-2): Table memo translation" in text
    assert "![Figure 1](assets/23501/5/tiny.png)" in text


def test_docx_export_service_uses_simple_table_path_for_non_merged_cells(tmp_path: Path) -> None:
    export_dir = tmp_path / "exports"
    service = DocxExportService(export_dir=export_dir, project_root=tmp_path)
    simple_cells = [
        [
            {"text": "H1", "rowspan": 1, "colspan": 1, "header": True},
            {"text": "H2", "rowspan": 1, "colspan": 1, "header": True},
        ],
        [
            {"text": "V1", "rowspan": 1, "colspan": 1, "header": False},
            {"text": "V2", "rowspan": 1, "colspan": 1, "header": False},
        ],
    ]

    assert service._extract_simple_rows_from_cells(simple_cells) == [["H1", "H2"], ["V1", "V2"]]


def test_docx_export_service_includes_notes_and_images(tmp_path: Path) -> None:
    export_dir = tmp_path / "exports"
    media_dir = tmp_path / "artifacts" / "clause_browser_media" / "23501" / "5"
    media_dir.mkdir(parents=True, exist_ok=True)
    image_path = media_dir / "tiny.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wn9l5QAAAAASUVORK5CYII="
        )
    )
    service = DocxExportService(export_dir=export_dir, project_root=tmp_path)
    roots = [
        {
            "key": "23501:5",
            "clauseId": "5",
            "clauseTitle": "Session management",
            "text": "",
            "blocks": [
                {
                    "type": "paragraph",
                    "text": "Paragraph body",
                    "format": {"leftIndentPt": 18.0, "textIndentPt": -18.0},
                },
                {"type": "paragraph", "text": "Figure 5-1: Example figure caption", "format": {"styleName": "TF"}},
                {
                    "type": "table",
                    "cells": [
                        [{"text": "H1", "rowspan": 1, "colspan": 1, "header": True}],
                        [{"text": "V1", "rowspan": 1, "colspan": 1, "header": False}],
                    ],
                },
                {"type": "image", "src": "/clause-browser-media/23501/5/tiny.png", "alt": "Figure 1"},
            ],
            "children": [],
        }
    ]
    notes = [
        {
            "id": "23501:5:clause",
            "type": "clause",
            "clauseKey": "23501:5",
            "translation": "Clause memo translation",
            "sourceText": "Source clause text",
        },
        {
            "id": "selection:1",
            "type": "selection",
            "clauseKey": "23501:5",
            "blockIndex": 0,
            "translation": "Selection memo translation",
            "sourceText": "Selected source text",
        },
    ]

    result = service.export("My Export", roots, notes=notes)

    exported = Document(export_dir / result.file_name)
    texts = [paragraph.text for paragraph in exported.paragraphs]
    assert "Paragraph body" in texts
    caption = next(item for item in exported.paragraphs if item.text == "Figure 5-1: Example figure caption")
    assert caption.alignment == WD_ALIGN_PARAGRAPH.CENTER
    paragraph = next(item for item in exported.paragraphs if item.text == "Paragraph body")
    assert round(paragraph.paragraph_format.left_indent.pt, 1) == 18.0
    assert round(paragraph.paragraph_format.first_line_indent.pt, 1) == -18.0
    assert len(exported.inline_shapes) == 1
    assert len(exported.tables) == 1
    assert exported.tables[0].style.name == "Table Grid"


def test_docx_export_service_prefers_prebuilt_png_for_svg_images(tmp_path: Path) -> None:
    export_dir = tmp_path / "exports"
    media_dir = tmp_path / "artifacts" / "clause_browser_media" / "23501" / "5"
    media_dir.mkdir(parents=True, exist_ok=True)
    svg_path = media_dir / "diagram.svg"
    svg_path.write_text('<svg xmlns="http://www.w3.org/2000/svg"></svg>', encoding="utf-8")
    png_path = media_dir / "diagram.export.png"
    png_path.write_bytes(
        base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wn9l5QAAAAASUVORK5CYII=")
    )
    service = DocxExportService(export_dir=export_dir, project_root=tmp_path)

    resolved = service._prepare_export_image(svg_path)

    assert resolved == png_path


def test_docx_export_service_applies_paragraph_and_table_highlights(tmp_path: Path) -> None:
    export_dir = tmp_path / "exports"
    service = DocxExportService(export_dir=export_dir, project_root=tmp_path)
    roots = [
        {
            "key": "23501:5",
            "clauseId": "5",
            "clauseTitle": "Session management",
            "text": "",
            "blocks": [
                {
                    "type": "paragraph",
                    "text": "Paragraph body",
                    "format": {},
                },
                {
                    "type": "table",
                    "cells": [
                        [
                            {"text": "H1", "rowspan": 1, "colspan": 1, "header": True},
                            {"text": "H2", "rowspan": 1, "colspan": 1, "header": True},
                        ],
                        [
                            {"text": "V1", "rowspan": 1, "colspan": 1, "header": False},
                            {"text": "V2", "rowspan": 1, "colspan": 1, "header": False},
                        ],
                    ],
                },
            ],
            "children": [],
        }
    ]
    highlights = [
        {
            "id": "manual:paragraph",
            "clauseKey": "23501:5",
            "blockIndex": 0,
            "rowIndex": -1,
            "cellIndex": -1,
            "rowText": "Paragraph body",
        },
        {
            "id": "manual:cell",
            "clauseKey": "23501:5",
            "blockIndex": 1,
            "rowIndex": 1,
            "cellIndex": 1,
            "rowText": "V1 | V2",
        },
    ]

    result = service.export("My Export", roots, highlights=highlights)

    exported = Document(export_dir / result.file_name)
    paragraph = next(item for item in exported.paragraphs if item.text == "Paragraph body")
    assert 'w:highlight w:val="yellow"' in paragraph._p.xml
    highlighted_cell = exported.tables[0].cell(1, 1)
    plain_cell = exported.tables[0].cell(1, 0)
    assert 'w:fill="FFF59D"' in highlighted_cell._tc.xml
    assert 'w:fill="FFF59D"' not in plain_cell._tc.xml


def test_docx_export_service_uses_compact_mode_for_large_tables(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("SPECBOT_DOCX_MAX_PRECISE_TABLE_ROWS", "3")
    monkeypatch.setenv("SPECBOT_DOCX_MAX_PRECISE_TABLE_CELLS", "6")
    export_dir = tmp_path / "exports"
    service = DocxExportService(export_dir=export_dir, project_root=tmp_path)
    roots = [
        {
            "key": "23501:5",
            "clauseId": "5",
            "clauseTitle": "Session management",
            "text": "",
            "blocks": [
                {
                    "type": "table",
                    "rows": [
                        ["H1", "H2"],
                        ["V1", "V2"],
                        ["V3", "V4"],
                        ["V5", "V6"],
                    ],
                },
            ],
            "children": [],
        }
    ]
    highlights = [
        {
            "id": "manual:cell",
            "clauseKey": "23501:5",
            "blockIndex": 0,
            "rowIndex": 2,
            "cellIndex": 1,
            "rowText": "V3 | V4",
        },
    ]
    notes = [
        {
            "id": "selection:table-row",
            "type": "selection",
            "clauseKey": "23501:5",
            "blockIndex": 0,
            "rowIndex": 1,
            "translation": "Row memo",
            "sourceText": "V1 | V2",
        },
    ]

    result = service.export("My Export", roots, notes=notes, highlights=highlights)

    exported = Document(export_dir / result.file_name)
    assert len(exported.tables) == 0
    row_paragraph = next(item for item in exported.paragraphs if item.text == "V3 | V4")
    assert 'w:highlight w:val="yellow"' in row_paragraph._p.xml
    noted_paragraph = next(item for item in exported.paragraphs if item.text == "V1 | V2")
    assert "commentRangeStart" in noted_paragraph._p.xml


def test_docx_export_service_preserves_large_tables_when_requested(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("SPECBOT_DOCX_MAX_PRECISE_TABLE_ROWS", "3")
    monkeypatch.setenv("SPECBOT_DOCX_MAX_PRECISE_TABLE_CELLS", "6")
    export_dir = tmp_path / "exports"
    service = DocxExportService(export_dir=export_dir, project_root=tmp_path)
    roots = [
        {
            "key": "23501:5",
            "clauseId": "5",
            "clauseTitle": "Session management",
            "text": "",
            "blocks": [
                {
                    "type": "table",
                    "rows": [
                        ["H1", "H2"],
                        ["V1", "V2"],
                        ["V3", "V4"],
                        ["V5", "V6"],
                    ],
                },
            ],
            "children": [],
        }
    ]

    result = service.export("My Export", roots, preserve_large_tables=True)

    exported = Document(export_dir / result.file_name)
    assert len(exported.tables) == 1
    assert exported.tables[0].cell(3, 1).text == "V6"


def test_docx_export_service_simplifies_large_preserved_table_annotations_to_row_level(
    tmp_path: Path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("SPECBOT_DOCX_MAX_PRECISE_TABLE_ROWS", "3")
    monkeypatch.setenv("SPECBOT_DOCX_MAX_PRECISE_TABLE_CELLS", "6")
    export_dir = tmp_path / "exports"
    service = DocxExportService(export_dir=export_dir, project_root=tmp_path)
    roots = [
        {
            "key": "23501:5",
            "clauseId": "5",
            "clauseTitle": "Session management",
            "text": "",
            "blocks": [
                {
                    "type": "table",
                    "rows": [
                        ["H1", "H2"],
                        ["V1", "V2"],
                        ["V3", "V4"],
                        ["V5", "V6"],
                    ],
                },
            ],
            "children": [],
        }
    ]
    notes = [
        {
            "id": "selection:table-cell-1",
            "type": "selection",
            "clauseKey": "23501:5",
            "blockIndex": 0,
            "rowIndex": 2,
            "cellIndex": 0,
            "translation": "Left cell memo",
            "sourceText": "V3",
        },
        {
            "id": "selection:table-cell-2",
            "type": "selection",
            "clauseKey": "23501:5",
            "blockIndex": 0,
            "rowIndex": 2,
            "cellIndex": 1,
            "translation": "Right cell memo",
            "sourceText": "V4",
        },
    ]
    highlights = [
        {
            "id": "manual:cell",
            "clauseKey": "23501:5",
            "blockIndex": 0,
            "rowIndex": 2,
            "cellIndex": 1,
            "rowText": "V3 | V4",
        },
    ]

    result = service.export("My Export", roots, notes=notes, highlights=highlights, preserve_large_tables=True)

    exported_path = export_dir / result.file_name
    exported = Document(exported_path)
    assert len(exported.tables) == 1
    table = exported.tables[0]
    assert 'w:fill="FFF59D"' in table.cell(2, 0)._tc.xml
    assert 'w:fill="FFF59D"' in table.cell(2, 1)._tc.xml
    assert "commentRangeStart" in table.cell(2, 0).paragraphs[0]._p.xml
    assert "commentRangeStart" not in table.cell(2, 1).paragraphs[0]._p.xml

    with zipfile.ZipFile(exported_path) as archive:
        comments_xml = archive.read("word/comments.xml").decode("utf-8")
    assert "Cell 1: Left cell memo" in comments_xml
    assert "Cell 2: Right cell memo" in comments_xml


def test_docx_export_service_keeps_small_preserved_table_cell_level_annotations(
    tmp_path: Path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("SPECBOT_DOCX_MAX_PRECISE_TABLE_ROWS", "10")
    monkeypatch.setenv("SPECBOT_DOCX_MAX_PRECISE_TABLE_CELLS", "100")
    export_dir = tmp_path / "exports"
    service = DocxExportService(export_dir=export_dir, project_root=tmp_path)
    roots = [
        {
            "key": "23501:5",
            "clauseId": "5",
            "clauseTitle": "Session management",
            "text": "",
            "blocks": [
                {
                    "type": "table",
                    "rows": [
                        ["H1", "H2"],
                        ["V1", "V2"],
                    ],
                },
            ],
            "children": [],
        }
    ]
    notes = [
        {
            "id": "selection:table-cell",
            "type": "selection",
            "clauseKey": "23501:5",
            "blockIndex": 0,
            "rowIndex": 1,
            "cellIndex": 1,
            "translation": "Right cell memo",
            "sourceText": "V2",
        },
    ]
    highlights = [
        {
            "id": "manual:cell",
            "clauseKey": "23501:5",
            "blockIndex": 0,
            "rowIndex": 1,
            "cellIndex": 1,
            "rowText": "V1 | V2",
        },
    ]

    result = service.export("My Export", roots, notes=notes, highlights=highlights, preserve_large_tables=True)

    exported = Document(export_dir / result.file_name)
    table = exported.tables[0]
    assert 'w:fill="FFF59D"' not in table.cell(1, 0)._tc.xml
    assert 'w:fill="FFF59D"' in table.cell(1, 1)._tc.xml
    assert "commentRangeStart" not in table.cell(1, 0).paragraphs[0]._p.xml
    assert "commentRangeStart" in table.cell(1, 1).paragraphs[0]._p.xml


def test_docx_export_service_sorts_clauses_by_clause_order(tmp_path: Path) -> None:
    export_dir = tmp_path / "exports"
    service = DocxExportService(export_dir=export_dir, project_root=tmp_path)
    roots = [
        {
            "key": "23501:5.10",
            "specNo": "23501",
            "specTitle": "System architecture",
            "clauseId": "5.10",
            "clauseTitle": "Tenth clause",
            "clausePath": ["5", "5.10"],
            "orderInSource": 10,
            "blocks": [],
            "children": [],
        },
        {
            "key": "23501:5",
            "specNo": "23501",
            "specTitle": "System architecture",
            "clauseId": "5",
            "clauseTitle": "Session management",
            "clausePath": ["5"],
            "orderInSource": 1,
            "blocks": [],
            "children": [
                {
                    "key": "23501:5.2",
                    "specNo": "23501",
                    "specTitle": "System architecture",
                    "clauseId": "5.2",
                    "clauseTitle": "Second child",
                    "clausePath": ["5", "5.2"],
                    "orderInSource": 3,
                    "blocks": [],
                    "children": [],
                },
                {
                    "key": "23501:5.1",
                    "specNo": "23501",
                    "specTitle": "System architecture",
                    "clauseId": "5.1",
                    "clauseTitle": "First child",
                    "clausePath": ["5", "5.1"],
                    "orderInSource": 2,
                    "blocks": [],
                    "children": [],
                },
            ],
        },
    ]

    result = service.export("My Export", roots)

    exported = Document(export_dir / result.file_name)
    headings = [p.text for p in exported.paragraphs if p.style.name.startswith("Heading")]
    assert headings == [
        "23501 System architecture",
        "5 Session management",
        "5.1 First child",
        "5.2 Second child",
        "5.10 Tenth clause",
    ]


def test_sanitize_file_stem_removes_invalid_characters() -> None:
    assert sanitize_file_stem('A/B:C*D? E') == "ABCD_E"
