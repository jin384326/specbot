from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.simpletypes import ST_Merge
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

from parser.models import ClauseDoc, DocRecord, PassageDoc, TableDoc, TableRowDoc

CLAUSE_PATTERN = re.compile(
    r"^(?P<clause>(?:\d+(?:[A-Za-z])?(?:\.\d+(?:[A-Za-z])?)*)|(?:[A-Z](?:\.\d+(?:[A-Za-z])?)+))\s+(?P<title>.+)$"
)
ANNEX_PATTERN = re.compile(r"^(?P<clause>Annex\s+[A-Z])(?:\s*\([^)]+\))?\s+(?P<title>.+)$")
SPEC_REF_PATTERN = re.compile(r"\b(?:3GPP\s+)?(?:TS|TR)\s+(\d{2}\.\d{3}|\d{5})\b")
STAGE_PATTERN = re.compile(r"\b(Stage\s+\d+)\b", re.IGNORECASE)


@dataclass
class SpecMetadata:
    spec_no: str = ""
    spec_title: str = ""
    release: str = ""
    release_data: str = ""
    series: str = ""
    ts_or_tr: str = ""
    stage_hint: str = ""
    version_tag: str = ""
    source_file: str = ""


@dataclass
class ClauseBuffer:
    clause_id: str
    clause_title: str
    clause_path: list[str]
    parent_clause_id: str
    level: int
    paragraphs: list[str] = field(default_factory=list)
    paragraph_indices: list[int] = field(default_factory=list)
    order_in_source: int = 0
    excluded: bool = False


def iter_block_items(parent: DocxDocument | _Cell) -> Iterator[Paragraph | Table]:
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f"Unsupported parent type: {type(parent)!r}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def normalize_whitespace(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def text_fingerprint(value: str) -> str:
    """Ignore all whitespace for equality (Word line breaks inside words → spurious spaces)."""
    return re.sub(r"\s+", "", (value or "").lower())


BRACKET_SEGMENT_PATTERN = re.compile(r"\[([^\]]*)\]")


def dedupe_duplicate_brackets(text: str) -> str:
    """Drop repeated [...] blocks that have the same inner text (case-insensitive, whitespace-normalized)."""
    seen: set[str] = set()
    parts: list[str] = []
    pos = 0
    for match in BRACKET_SEGMENT_PATTERN.finditer(text):
        parts.append(text[pos : match.start()])
        inner_key = text_fingerprint(match.group(1))
        if inner_key in seen:
            pos = match.end()
            continue
        seen.add(inner_key)
        parts.append(match.group(0))
        pos = match.end()
    parts.append(text[pos:])
    return normalize_whitespace("".join(parts))


def remove_redundant_brackets_matching_outside(text: str) -> str:
    """Remove [inner] when the same text already appears outside brackets in the cell (typical spec table duplication)."""
    stripped = BRACKET_SEGMENT_PATTERN.sub(" ", text)
    stripped_norm = normalize_whitespace(stripped).lower()
    if not stripped_norm:
        return text
    parts: list[str] = []
    pos = 0
    for match in BRACKET_SEGMENT_PATTERN.finditer(text):
        parts.append(text[pos : match.start()])
        inner = normalize_whitespace(match.group(1))
        key = inner.lower()
        if len(inner) >= 3 and key and key in stripped_norm:
            pos = match.end()
            continue
        parts.append(match.group(0))
        pos = match.end()
    parts.append(text[pos:])
    return normalize_whitespace("".join(parts))


def _dedupe_semicolon_list_items_in_line(line: str, *, min_segment_len: int = 14) -> str:
    """Drop repeated fragments in 'a; b; a' style lists inside one line (common in spec table cells)."""
    parts = re.split(r"\s*;\s*", line)
    out: list[str] = []
    seen: set[str] = set()
    prev_fp: str | None = None
    for part in parts:
        p = normalize_whitespace(part)
        if not p:
            continue
        fp = text_fingerprint(p)
        if prev_fp is not None and fp == prev_fp:
            continue
        if len(p) >= min_segment_len and fp in seen:
            continue
        if len(p) >= min_segment_len:
            seen.add(fp)
        out.append(p)
        prev_fp = fp
    return "; ".join(out)


def dedupe_repeated_lines_and_semicolon_lists(text: str, *, min_line_len: int = 20) -> str:
    """Remove duplicate lines and semicolon-separated duplicates (list-like cell content)."""
    lines = text.split("\n")
    result_lines: list[str] = []
    seen_fp: set[str] = set()
    prev_line_fp: str | None = None
    for raw_line in lines:
        line = _dedupe_semicolon_list_items_in_line(raw_line)
        line = normalize_whitespace(line)
        if not line:
            continue
        fp = text_fingerprint(line)
        if fp == prev_line_fp:
            continue
        if len(line) >= min_line_len and fp in seen_fp:
            continue
        if len(line) >= min_line_len:
            seen_fp.add(fp)
        result_lines.append(line)
        prev_line_fp = fp
    return "\n".join(result_lines)


def normalize_table_cell_text(text: str) -> str:
    """Normalize a single table cell: whitespace, dedupe bracket groups, drop brackets redundant with surrounding text."""
    raw = normalize_whitespace(text or "")
    if not raw:
        return ""
    raw = dedupe_duplicate_brackets(raw)
    raw = remove_redundant_brackets_matching_outside(raw)
    raw = dedupe_repeated_lines_and_semicolon_lists(raw)
    return normalize_whitespace(raw)


def dedupe_consecutive_duplicate_paragraphs(
    paragraphs: list[str],
    indices: list[int] | None = None,
) -> tuple[list[str], list[int]]:
    """Drop back-to-back identical paragraphs (e.g. duplicated list items in Word).

    If indices is provided (same length as paragraphs), keeps the index aligned with kept paragraphs.
    """
    out_p: list[str] = []
    out_i: list[int] = []
    prev_key: str | None = None
    for idx, paragraph in enumerate(paragraphs):
        fp = text_fingerprint(paragraph)
        if not fp:
            continue
        if fp == prev_key:
            continue
        out_p.append(paragraph)
        out_i.append(indices[idx] if indices is not None and idx < len(indices) else idx + 1)
        prev_key = fp
    return out_p, out_i


def linearized_row_pairs(header: list[str], row: list[str]) -> list[tuple[str, str]]:
    """Build (label, cell) pairs; skip consecutive duplicate cell values (horizontal merged cells in Word)."""
    pairs: list[tuple[str, str]] = []
    prev_fp: str | None = None
    for col_idx, cell in enumerate(row):
        label = header[col_idx] if col_idx < len(header) and header[col_idx] else f"column_{col_idx + 1}"
        val = normalize_whitespace(cell) if cell else ""
        if not val:
            continue
        fp = text_fingerprint(val)
        if prev_fp is not None and fp == prev_fp:
            continue
        pairs.append((label, val))
        prev_fp = fp
    return pairs


def dedupe_duplicate_cell_texts_preserve_order(
    row: list[str],
    *,
    drop_empty: bool = True,
) -> list[str]:
    """Remove duplicate cell strings in row order (Word merge often repeats the same label with empty cells between).

    Uses text_fingerprint so spacing/unicode variants still match. Empty cells are omitted from the result when
    drop_empty is True (cleaner JSON lists).
    """
    seen: set[str] = set()
    out: list[str] = []
    for cell in row:
        raw = cell or ""
        val = normalize_whitespace(raw)
        if not val:
            if not drop_empty:
                out.append(raw)
            continue
        fp = text_fingerprint(val)
        if fp in seen:
            continue
        seen.add(fp)
        out.append(raw)
    return out


def paragraph_style_level(style_name: str) -> int | None:
    match = re.match(r"Heading\s+(\d+)$", style_name or "")
    return int(match.group(1)) if match else None


def split_clause_heading(text: str) -> tuple[str, str] | None:
    normalized = normalize_whitespace(text)
    annex_match = ANNEX_PATTERN.match(normalized)
    if annex_match:
        return annex_match.group("clause"), annex_match.group("title")
    match = CLAUSE_PATTERN.match(normalized)
    if not match:
        return None
    clause_id = match.group("clause")
    if not is_probable_clause_id(clause_id):
        return None
    return clause_id, match.group("title")


def should_treat_paragraph_as_heading(
    text: str,
    heading_level: int | None,
    heading: tuple[str, str] | None,
) -> bool:
    if heading_level is not None:
        return True
    if heading is None:
        return False
    clause_id, _clause_title = heading
    if clause_id.startswith("Annex "):
        return False
    if text.endswith("."):
        return False
    return True


def is_probable_clause_id(clause_id: str) -> bool:
    if clause_id.startswith("Annex "):
        return True
    if clause_id.isdigit():
        return True
    if "." in clause_id:
        return True
    if clause_id[:-1].isdigit() and clause_id[-1].islower():
        return True
    return False


def _flatten_table_cells_row_major(table: Table) -> list[_Cell]:
    """Expand the table to one ``_Cell`` per layout-grid slot (same as ``Table._cells``).

    Some 3GPP (and other) documents contain invalid vertical-merge markup where
    ``w:vMerge w:val=\"continue\"`` appears before a full first row exists in the
    grid; ``Table._cells`` then does ``cells[-col_count]`` and raises ``IndexError``.
    In that case we fall back to a real ``w:tc`` cell so parsing can continue.
    """
    col_count = table._column_count
    cells: list[_Cell] = []
    for tc in table._tbl.iter_tcs():
        for grid_span_idx in range(tc.grid_span):
            if tc.vMerge == ST_Merge.CONTINUE:
                if len(cells) >= col_count:
                    cells.append(cells[-col_count])
                else:
                    cells.append(_Cell(tc, table))
            elif grid_span_idx > 0:
                if cells:
                    cells.append(cells[-1])
                else:
                    cells.append(_Cell(tc, table))
            else:
                cells.append(_Cell(tc, table))
    return cells


def clean_table_matrix(table: Table) -> list[list[str]]:
    """Build a logical matrix from a row-major cell grid (not ``row.cells``).

    ``row.cells`` builds fresh ``_Cell`` wrappers per access; ``id(cell._tc)`` can
    collide across rows in CPython, breaking merge detection. A flattened grid repeats
    the same ``_Cell`` instance for horizontal spans and vertical continuations, so
    ``id(cell)`` identifies merged regions. We use :func:`_flatten_table_cells_row_major`
    instead of ``Table._cells`` to survive malformed ``vMerge`` in the wild.
    """
    col_count = table._column_count
    flat = _flatten_table_cells_row_major(table)
    if col_count == 0 or not flat:
        return []
    nrows = len(flat) // col_count
    matrix: list[list[str]] = []
    emitted_cell_ids: set[int] = set()
    for row_idx in range(nrows):
        start = row_idx * col_count
        row_cells = flat[start : start + col_count]
        row_out: list[str] = []
        prev_cell: _Cell | None = None
        for cell in row_cells:
            if prev_cell is not None and cell is prev_cell:
                continue
            prev_cell = cell
            cid = id(cell)
            if cid in emitted_cell_ids:
                row_out.append("")
            else:
                row_out.append(normalize_table_cell_text(cell.text))
                emitted_cell_ids.add(cid)
        if any(normalize_whitespace(c) for c in row_out):
            matrix.append(row_out)
    return matrix


def table_to_markdown(matrix: list[list[str]]) -> str:
    if not matrix:
        return ""
    width = max(len(row) for row in matrix)
    padded = [row + [""] * (width - len(row)) for row in matrix]
    header = padded[0]
    separator = ["---"] * width
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def table_to_linearized_text(matrix: list[list[str]], table_title: str = "") -> str:
    parts: list[str] = []
    if table_title:
        parts.append(f"Table {table_title}")
    if not matrix:
        return ""
    header = matrix[0]
    for idx, row in enumerate(matrix[1:], start=1):
        pairs = [f"{label}: {cell}" for label, cell in linearized_row_pairs(header, row)]
        if pairs:
            parts.append(f"row {idx}: " + "; ".join(pairs))
    if len(matrix) == 1:
        parts.append("header: " + "; ".join(cell for cell in header if cell))
    return "\n".join(parts)


def referenced_specs_from_text(text: str) -> list[str]:
    return sorted({match.group(1).replace(".", "") for match in SPEC_REF_PATTERN.finditer(text)})


def is_excluded_clause(clause_id: str, clause_title: str) -> bool:
    normalized_title = normalize_whitespace(clause_title).lower()
    return clause_id.startswith("Annex ") or "change history" in normalized_title


def infer_spec_title(document: DocxDocument) -> str:
    core_title = normalize_whitespace(document.core_properties.title or "")
    if document.tables:
        for row in document.tables[0].rows:
            cells = [normalize_whitespace(cell.text) for cell in row.cells]
            row_text = next((cell for cell in cells if cell), "")
            if not row_text or "3rd Generation Partnership Project" not in row_text:
                continue
            parts = [part.strip() for part in row_text.split(";") if part.strip()]
            if len(parts) >= 2:
                if "stage" in parts[-1].lower() or "release" in parts[-1].lower():
                    return "; ".join(parts[-2:])
                return parts[-1]
            return row_text
    if core_title and not re.match(r"^3GPP\s+(?:TS|TR)\s+\d{2}\.\d{3}$", core_title):
        return core_title
    for paragraph in document.paragraphs:
        text = normalize_whitespace(paragraph.text)
        if text and not paragraph.style.name.lower().startswith("toc"):
            return text
    return ""


def infer_stage_hint(spec_title: str, document: DocxDocument) -> str:
    title_match = STAGE_PATTERN.search(spec_title)
    if title_match:
        return title_match.group(1).title()
    if document.tables:
        for row in document.tables[0].rows:
            row_text = " ".join(normalize_whitespace(cell.text) for cell in row.cells if normalize_whitespace(cell.text))
            match = STAGE_PATTERN.search(row_text)
            if match:
                return match.group(1).title()
    return "else"


class DocxClauseParser:
    def __init__(
        self,
        passage_char_limit: int = 1200,
        passage_paragraph_limit: int = 4,
        min_passage_chars: int = 350,
    ) -> None:
        self.passage_char_limit = passage_char_limit
        self.passage_paragraph_limit = passage_paragraph_limit
        self.min_passage_chars = min_passage_chars

    def parse(self, docx_path: str | Path, metadata: SpecMetadata | dict[str, Any] | None = None) -> list[DocRecord]:
        path = Path(docx_path)
        document = Document(str(path))
        spec_metadata = self._coerce_metadata(path, metadata, document)
        records: list[DocRecord] = []
        clause_stack: list[ClauseBuffer] = []
        active_clause: ClauseBuffer | None = None
        fallback_clause_counter = 0
        paragraph_counter = 0
        table_counter = 0
        clause_counter = 0
        last_text_title = ""

        for block in iter_block_items(document):
            if isinstance(block, Paragraph):
                text = normalize_whitespace(block.text)
                style_name = block.style.name if block.style else ""
                if not text or style_name.lower().startswith("toc"):
                    continue

                heading_level = paragraph_style_level(style_name)
                heading = split_clause_heading(text)
                if not should_treat_paragraph_as_heading(text, heading_level, heading):
                    heading = None
                if heading_level is not None or heading is not None:
                    if active_clause is not None:
                        records.extend(self._emit_clause_records(active_clause, spec_metadata))
                    clause_counter += 1
                    paragraph_counter += 1
                    clause_id, clause_title, level = self._resolve_heading(text, heading_level, heading, clause_counter)
                    if clause_id and clause_id[0].isdigit():
                        while clause_stack and clause_stack[-1].clause_id.startswith("Annex "):
                            clause_stack.pop()
                    while clause_stack and clause_stack[-1].level >= level:
                        clause_stack.pop()
                    clause_path = [*clause_stack[-1].clause_path, clause_id] if clause_stack else [clause_id]
                    parent_clause_id = clause_stack[-1].clause_id if clause_stack else ""
                    active_clause = ClauseBuffer(
                        clause_id=clause_id,
                        clause_title=clause_title,
                        clause_path=clause_path,
                        parent_clause_id=parent_clause_id,
                        level=level,
                        order_in_source=paragraph_counter,
                        excluded=is_excluded_clause(clause_id, clause_title),
                    )
                    clause_stack.append(active_clause)
                    last_text_title = clause_title
                    continue

                paragraph_counter += 1
                if active_clause is None:
                    fallback_clause_counter += 1
                    fallback_id = f"front_matter_{fallback_clause_counter}"
                    active_clause = ClauseBuffer(
                        clause_id=fallback_id,
                        clause_title="Front matter",
                        clause_path=[fallback_id],
                        parent_clause_id="",
                        level=1,
                        order_in_source=paragraph_counter,
                    )
                    clause_stack = [active_clause]
                active_clause.paragraphs.append(text)
                active_clause.paragraph_indices.append(paragraph_counter)
                if not last_text_title:
                    last_text_title = active_clause.clause_title
            else:
                matrix = clean_table_matrix(block)
                if not matrix:
                    continue
                table_counter += 1
                table_id = f"{spec_metadata.spec_no or path.stem}:table:{table_counter}"
                table_title = last_text_title or (active_clause.clause_title if active_clause else f"Table {table_counter}")
                if (active_clause and active_clause.excluded) or "change history" in normalize_whitespace(table_title).lower():
                    continue
                linearized_text = table_to_linearized_text(matrix, table_title)
                common_fields = self._common_fields(spec_metadata, active_clause, path, table_counter)
                header = matrix[0] if matrix else []
                header_compact = dedupe_duplicate_cell_texts_preserve_order(header)
                table_doc = TableDoc(
                    **common_fields,
                    doc_id=table_id,
                    text=linearized_text,
                    table_id=table_id,
                    table_title=table_title,
                    table_raw=matrix,
                    table_markdown=table_to_markdown(matrix),
                    table_headers=header_compact,
                    referenced_specs=referenced_specs_from_text(linearized_text),
                )
                records.append(table_doc)
                for row_index, row in enumerate(matrix[1:], start=1):
                    row_header = row[0] if row else ""
                    row_cells_compact = dedupe_duplicate_cell_texts_preserve_order(row)
                    row_text = "; ".join(f"{label}: {cell}" for label, cell in linearized_row_pairs(header, row))
                    row_doc = TableRowDoc(
                        **common_fields,
                        doc_id=f"{table_id}:row:{row_index}",
                        text=row_text,
                        table_id=table_id,
                        table_title=table_title,
                        table_raw=matrix,
                        table_markdown=table_doc.table_markdown,
                        row_index=row_index,
                        row_header=row_header,
                        row_cells=row_cells_compact,
                        referenced_specs=referenced_specs_from_text(row_text),
                    )
                    records.append(row_doc)

        if active_clause is not None:
            records.extend(self._emit_clause_records(active_clause, spec_metadata))
        return records

    def _coerce_metadata(
        self,
        path: Path,
        metadata: SpecMetadata | dict[str, Any] | None,
        document: DocxDocument,
    ) -> SpecMetadata:
        if isinstance(metadata, SpecMetadata):
            result = metadata
        else:
            result = SpecMetadata(**(metadata or {}))
        if not result.source_file:
            result.source_file = str(path)
        if not result.spec_no:
            digits = re.match(r"(?P<spec>\d{5})", path.stem)
            if digits:
                result.spec_no = digits.group("spec")
                result.series = result.series or digits.group("spec")[:2]
        if not result.release:
            parts = path.parts
            for part in parts:
                if part.startswith("Rel-"):
                    result.release = part
        if not result.release_data:
            dated_path = path.as_posix()
            date_match = re.search(r"/(\d{4}-\d{2})/", dated_path)
            if date_match:
                result.release_data = date_match.group(1)
        if not result.version_tag:
            version = re.search(r"-([a-z]\d+)$", path.stem, re.IGNORECASE)
            if version:
                result.version_tag = version.group(1)
        if result.spec_no and not result.ts_or_tr:
            result.ts_or_tr = "TS"
        if not result.spec_title:
            result.spec_title = infer_spec_title(document)
        if not result.stage_hint:
            result.stage_hint = infer_stage_hint(result.spec_title, document)
        return result

    def _resolve_heading(
        self,
        text: str,
        heading_level: int | None,
        heading: tuple[str, str] | None,
        clause_counter: int,
    ) -> tuple[str, str, int]:
        if heading:
            clause_id, clause_title = heading
            if clause_id.startswith("Annex "):
                level = 1
            else:
                level = clause_id.count(".") + 1
            return clause_id, clause_title, level
        clause_id = f"heading_{clause_counter}"
        clause_title = text
        level = heading_level or 1
        return clause_id, clause_title, level

    def _common_fields(
        self,
        metadata: SpecMetadata,
        clause: ClauseBuffer | None,
        path: Path,
        order_in_source: int,
    ) -> dict[str, Any]:
        clause_id = clause.clause_id if clause else ""
        clause_title = clause.clause_title if clause else ""
        clause_path = clause.clause_path if clause else []
        parent_clause_id = clause.parent_clause_id if clause else ""
        return {
            "spec_no": metadata.spec_no,
            "spec_title": metadata.spec_title,
            "release": metadata.release,
            "release_data": metadata.release_data,
            "series": metadata.series,
            "ts_or_tr": metadata.ts_or_tr,
            "stage_hint": metadata.stage_hint,
            "clause_id": clause_id,
            "clause_title": clause_title,
            "clause_path": clause_path,
            "parent_clause_id": parent_clause_id,
            "source_file": metadata.source_file or str(path),
            "order_in_source": order_in_source,
            "version_tag": metadata.version_tag,
        }

    def _emit_clause_records(self, clause: ClauseBuffer, metadata: SpecMetadata) -> list[DocRecord]:
        if clause.excluded:
            return []
        paragraphs, paragraph_indices = dedupe_consecutive_duplicate_paragraphs(
            clause.paragraphs,
            clause.paragraph_indices if len(clause.paragraph_indices) == len(clause.paragraphs) else None,
        )
        text = "\n".join(paragraphs).strip()
        if not text:
            return []
        clause_id_prefix = metadata.spec_no or Path(metadata.source_file).stem
        common_fields = self._common_fields(metadata, clause, Path(metadata.source_file), clause.order_in_source)
        clause_doc = ClauseDoc(
            **common_fields,
            doc_id=f"{clause_id_prefix}:clause:{clause.clause_id}",
            text=text,
            referenced_specs=referenced_specs_from_text(text),
        )
        records: list[DocRecord] = [clause_doc]
        records.extend(self._build_passages(clause_doc, paragraphs, paragraph_indices))
        return records

    def _build_passages(
        self,
        clause_doc: ClauseDoc,
        paragraphs: list[str],
        paragraph_indices: list[int],
    ) -> list[PassageDoc]:
        if len(paragraphs) <= 1 and len(clause_doc.text) < self.passage_char_limit:
            return []
        passages: list[PassageDoc] = []
        chunk: list[str] = []
        chunk_indices: list[int] = []
        for paragraph, idx in zip(paragraphs, paragraph_indices):
            candidate = "\n".join([*chunk, paragraph])
            if chunk and (
                len(candidate) > self.passage_char_limit or len(chunk) >= self.passage_paragraph_limit
            ):
                passages.append(self._make_passage(clause_doc, passages, chunk, chunk_indices))
                chunk = []
                chunk_indices = []
            chunk.append(paragraph)
            chunk_indices.append(idx)
        if chunk:
            passages.append(self._make_passage(clause_doc, passages, chunk, chunk_indices))
        if len(passages) == 1 and len(passages[0].text) < self.min_passage_chars:
            return []
        return passages

    def _make_passage(
        self,
        clause_doc: ClauseDoc,
        existing_passages: list[PassageDoc],
        paragraphs: list[str],
        paragraph_indices: list[int],
    ) -> PassageDoc:
        passage_index = len(existing_passages)
        text = "\n".join(paragraphs).strip()
        payload = clause_doc.to_dict()
        payload.update(
            {
                "doc_id": f"{clause_doc.doc_id}:passage:{passage_index}",
                "doc_type": "passage_doc",
                "content_kind": "passage",
                "text": text,
                "passage_id": f"{clause_doc.doc_id}:passage:{passage_index}",
                "passage_index": passage_index,
                "paragraph_start_index": paragraph_indices[0],
                "paragraph_end_index": paragraph_indices[-1],
                "referenced_specs": referenced_specs_from_text(text),
            }
        )
        return PassageDoc(**payload)
