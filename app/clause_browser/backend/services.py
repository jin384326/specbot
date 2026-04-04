from __future__ import annotations

import os
import re
import subprocess
import sys
import json
import asyncio
import threading
from io import BytesIO
from urllib import error as urllib_error
from urllib import request as urllib_request
from dataclasses import dataclass
from functools import cmp_to_key
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from docx.text.paragraph import Paragraph


INVALID_FILENAME_CHARS = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
WHITESPACE_RUN = re.compile(r"\s+")
TRANSLATION_CHUNK_LIMIT = 12000
CAPTION_PARAGRAPH = re.compile(r"^(Figure|Table)\s+[A-Za-z0-9.\-]+(?:\s*[:.\-]\s*|\s{2,})")
DOCX_HIGHLIGHT_COLOR = "yellow"
DOCX_SHADE_FILL = "FFF59D"


@dataclass(frozen=True)
class ExportResult:
    title: str
    file_name: str
    absolute_path: str
    relative_path: str
    clause_count: int

    def to_dict(self) -> dict[str, object]:
        return {
            "title": self.title,
            "fileName": self.file_name,
            "absolutePath": self.absolute_path,
            "relativePath": self.relative_path,
            "clauseCount": self.clause_count,
        }


class LLMActionQueueFullError(RuntimeError):
    pass


class LLMActionCancelledError(RuntimeError):
    pass


class DocxExportService:
    def __init__(self, export_dir: str | Path, project_root: str | Path) -> None:
        self._export_dir = Path(export_dir)
        self._project_root = Path(project_root)
        self._media_root = self._project_root / "artifacts" / "clause_browser_media"
        self._export_dir.mkdir(parents=True, exist_ok=True)

    def export(
        self,
        title: str,
        roots: list[dict[str, Any]],
        notes: list[dict[str, Any]] | None = None,
        highlights: list[dict[str, Any]] | None = None,
    ) -> ExportResult:
        cleaned_title, document, clause_count = self._build_document(
            title=title,
            roots=roots,
            notes=notes or [],
            highlights=highlights or [],
        )

        file_name = self._allocate_file_name(cleaned_title)
        absolute_path = self._export_dir / file_name
        document.save(absolute_path)
        try:
            relative_path = absolute_path.relative_to(self._project_root)
        except ValueError:
            relative_path = absolute_path

        return ExportResult(
            title=cleaned_title,
            file_name=file_name,
            absolute_path=str(absolute_path),
            relative_path=str(relative_path),
            clause_count=clause_count,
        )

    def export_bytes(
        self,
        title: str,
        roots: list[dict[str, Any]],
        notes: list[dict[str, Any]] | None = None,
        highlights: list[dict[str, Any]] | None = None,
    ) -> tuple[str, bytes]:
        cleaned_title, document, _clause_count = self._build_document(
            title=title,
            roots=roots,
            notes=notes or [],
            highlights=highlights or [],
        )
        file_name = f"{sanitize_file_stem(cleaned_title)}.docx"
        buffer = BytesIO()
        document.save(buffer)
        return file_name, buffer.getvalue()

    def _build_document(
        self,
        *,
        title: str,
        roots: list[dict[str, Any]],
        notes: list[dict[str, Any]],
        highlights: list[dict[str, Any]],
    ) -> tuple[str, Document, int]:
        cleaned_title = title.strip()
        if not cleaned_title:
            raise ValueError("Document title is required.")
        if not roots:
            raise ValueError("At least one loaded clause is required to export a DOCX.")

        document = Document()
        document.add_heading(cleaned_title, level=0)
        note_index = self._index_notes(notes)
        highlight_index = self._index_highlights(highlights)
        clause_count = 0
        grouped_roots: dict[str, list[dict[str, Any]]] = {}
        spec_titles: dict[str, str] = {}
        for root in roots:
            spec_no = str(root.get("specNo") or "").strip() or "Unknown Spec"
            if spec_no not in grouped_roots:
                grouped_roots[spec_no] = []
                spec_titles[spec_no] = str(root.get("specTitle") or "").strip()
            grouped_roots[spec_no].append(self._sort_clause_tree(root))

        for spec_no in sorted(grouped_roots.keys(), key=cmp_to_key(self._compare_mixed_token)):
            spec_title = spec_titles.get(spec_no, "")
            spec_heading = spec_no if not spec_title else f"{spec_no} {spec_title}"
            document.add_heading(spec_heading, level=1)
            for root in sorted(grouped_roots.get(spec_no, []), key=cmp_to_key(self._compare_clause_nodes)):
                clause_count += self._write_clause(
                    document,
                    root,
                    depth=1,
                    note_index=note_index,
                    highlight_index=highlight_index,
                )
        return cleaned_title, document, clause_count

    def _allocate_file_name(self, title: str) -> str:
        base_name = sanitize_file_stem(title)
        candidate = f"{base_name}.docx"
        suffix = 2
        while (self._export_dir / candidate).exists():
            candidate = f"{base_name}_{suffix}.docx"
            suffix += 1
        return candidate

    def _write_clause(
        self,
        document: Document,
        clause: dict[str, Any],
        depth: int,
        note_index: dict[str, Any],
        highlight_index: dict[tuple[str, int], list[dict[str, Any]]],
    ) -> int:
        clause_id = str(clause.get("clauseId") or "").strip()
        clause_title = str(clause.get("clauseTitle") or "").strip()
        clause_key = str(clause.get("key") or "").strip()
        heading_text = " ".join(part for part in [clause_id, clause_title] if part).strip() or clause_id or clause_title
        heading_paragraph = document.add_heading(heading_text, level=min(depth + 1, 9))

        blocks = list(clause.get("blocks") or [])
        if blocks:
            for block_index, block in enumerate(blocks):
                self._write_block(
                    document,
                    block,
                    clause_key=clause_key,
                    block_index=block_index,
                    note_index=note_index,
                    highlight_index=highlight_index,
                )
        else:
            body = str(clause.get("text") or "").strip()
            if body:
                for paragraph in body.splitlines():
                    trimmed = paragraph.strip()
                    if trimmed:
                        exported_paragraph = document.add_paragraph(trimmed)
                        if self._is_caption_paragraph(trimmed, {}):
                            exported_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for note in note_index.get("clause", {}).get(clause_key, []):
            self._attach_comment_to_paragraph(heading_paragraph, note)

        total = 1
        for child in clause.get("children") or []:
            total += self._write_clause(document, child, depth + 1, note_index=note_index, highlight_index=highlight_index)
        return total

    def _sort_clause_tree(self, clause: dict[str, Any]) -> dict[str, Any]:
        children = [self._sort_clause_tree(child) for child in clause.get("children") or []]
        sorted_children = sorted(children, key=cmp_to_key(self._compare_clause_nodes))
        return {**clause, "children": sorted_children}

    def _write_block(
        self,
        document: Document,
        block: dict[str, Any],
        *,
        clause_key: str,
        block_index: int,
        note_index: dict[str, Any],
        highlight_index: dict[tuple[str, int], list[dict[str, Any]]],
    ) -> None:
        block_type = str(block.get("type") or "")
        selection_notes = note_index.get("selection", {}).get((clause_key, block_index), [])
        block_highlights = highlight_index.get((clause_key, block_index), [])
        if block_type == "paragraph":
            text = str(block.get("text") or "").strip()
            if text:
                paragraph = self._add_paragraph_with_notes(document, text, selection_notes)
                if any(
                    int(item.get("rowIndex", -1)) < 0 and int(item.get("cellIndex", -1)) < 0
                    for item in block_highlights
                ):
                    self._apply_paragraph_highlight(paragraph)
                self._apply_paragraph_format(paragraph, block.get("format") or {})
                if self._is_caption_paragraph(text, block.get("format") or {}):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif block_type == "table":
            self._write_table(document, block, selection_notes, block_highlights)
        elif block_type == "image":
            self._write_image(document, block)

    def _write_table(
        self,
        document: Document,
        block: dict[str, Any],
        selection_notes: list[dict[str, Any]],
        block_highlights: list[dict[str, Any]],
    ) -> None:
        cells = list(block.get("cells") or [])
        if not cells:
            rows = list(block.get("rows") or [])
            if not rows:
                return
            table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
            table.style = "Table Grid"
            anchor_paragraphs: list[Paragraph] = []
            row_anchor_paragraphs: dict[int, list[Paragraph]] = {}
            row_cells: dict[int, list[Any]] = {}
            logical_cells: dict[tuple[int, int], Any] = {}
            for row_idx, row in enumerate(rows):
                for col_idx, value in enumerate(row):
                    cell = table.cell(row_idx, col_idx)
                    paragraph = cell.paragraphs[0]
                    paragraph.text = str(value or "")
                    anchor_paragraphs.append(paragraph)
                    row_anchor_paragraphs.setdefault(row_idx, []).append(paragraph)
                    row_cells.setdefault(row_idx, []).append(cell)
                    logical_cells[(row_idx, col_idx)] = cell
            self._apply_table_highlights(row_cells, logical_cells, block_highlights)
            self._attach_table_selection_notes(anchor_paragraphs, row_anchor_paragraphs, selection_notes)
            return

        row_count = len(cells)
        col_count = max(sum(int(cell.get("colspan") or 1) for cell in row) for row in cells)
        table = document.add_table(rows=row_count, cols=col_count)
        table.style = "Table Grid"
        occupied: set[tuple[int, int]] = set()
        anchor_paragraphs: list[Paragraph] = []
        row_anchor_paragraphs: dict[int, list[Paragraph]] = {}
        row_cells: dict[int, list[Any]] = {}
        logical_cells: dict[tuple[int, int], Any] = {}
        for row_idx, row in enumerate(cells):
            col_idx = 0
            for logical_idx, cell_data in enumerate(row):
                while (row_idx, col_idx) in occupied:
                    col_idx += 1
                rowspan = max(1, int(cell_data.get("rowspan") or 1))
                colspan = max(1, int(cell_data.get("colspan") or 1))
                base_cell = table.cell(row_idx, col_idx)
                paragraph = base_cell.paragraphs[0]
                paragraph.text = str(cell_data.get("text") or "")
                anchor_paragraphs.append(paragraph)
                row_anchor_paragraphs.setdefault(row_idx, []).append(paragraph)
                row_cells.setdefault(row_idx, []).append(base_cell)
                logical_cells[(row_idx, logical_idx)] = base_cell
                if colspan > 1:
                    base_cell = base_cell.merge(table.cell(row_idx, col_idx + colspan - 1))
                if rowspan > 1:
                    base_cell = base_cell.merge(table.cell(row_idx + rowspan - 1, col_idx + colspan - 1))
                for row_offset in range(rowspan):
                    for col_offset in range(colspan):
                        occupied.add((row_idx + row_offset, col_idx + col_offset))
                col_idx += colspan
        self._apply_table_highlights(row_cells, logical_cells, block_highlights)
        self._attach_table_selection_notes(anchor_paragraphs, row_anchor_paragraphs, selection_notes)

    def _apply_table_highlights(
        self,
        row_cells: dict[int, list[Any]],
        logical_cells: dict[tuple[int, int], Any],
        block_highlights: list[dict[str, Any]],
    ) -> None:
        for item in block_highlights:
            row_index = int(item.get("rowIndex", -1))
            cell_index = int(item.get("cellIndex", -1))
            if row_index < 0:
                continue
            if cell_index >= 0:
                cell = logical_cells.get((row_index, cell_index))
                if cell is not None:
                    self._apply_cell_highlight(cell)
                continue
            for cell in row_cells.get(row_index, []):
                self._apply_cell_highlight(cell)

    def _attach_table_selection_notes(
        self,
        anchor_paragraphs: list[Paragraph],
        row_anchor_paragraphs: dict[int, list[Paragraph]],
        selection_notes: list[dict[str, Any]],
    ) -> None:
        for note in selection_notes:
            row_index = int(note.get("rowIndex", -1) or -1)
            if row_index >= 0 and row_index in row_anchor_paragraphs:
                self._attach_comment_to_paragraphs(row_anchor_paragraphs[row_index], note)
            else:
                self._attach_comment_to_paragraphs(anchor_paragraphs, note)

    def _write_image(self, document: Document, block: dict[str, Any]) -> None:
        image_path = self._resolve_image_path(str(block.get("src") or ""))
        if image_path and image_path.exists():
            try:
                resolved_path = self._prepare_export_image(image_path)
                document.add_picture(str(resolved_path), width=Inches(6))
            except Exception:
                alt = str(block.get("alt") or "").strip()
                if alt:
                    document.add_paragraph(alt)
        else:
            alt = str(block.get("alt") or "").strip()
            if alt:
                document.add_paragraph(alt)

    def _resolve_image_path(self, src: str) -> Path | None:
        if not src:
            return None
        if src.startswith("/clause-browser-media/"):
            relative = src.removeprefix("/clause-browser-media/").lstrip("/")
            return self._media_root / relative
        candidate = Path(src)
        if candidate.is_absolute():
            return candidate
        return self._project_root / src.lstrip("./")

    def _prepare_export_image(self, path: Path) -> Path:
        if path.suffix.lower() != ".svg":
            return path
        png_path = path.with_suffix(".export.png")
        if png_path.exists():
            return png_path
        try:
            import cairosvg  # type: ignore
        except Exception as exc:
            raise RuntimeError(f"SVG export requires cairosvg: {exc}") from exc
        cairosvg.svg2png(url=str(path), write_to=str(png_path))
        return png_path

    @staticmethod
    def _apply_paragraph_format(paragraph, format_data: dict[str, Any]) -> None:
        left_indent_pt = format_data.get("leftIndentPt")
        text_indent_pt = format_data.get("textIndentPt")
        left_indent_px = format_data.get("leftIndentPx")
        text_indent_px = format_data.get("textIndentPx")
        if left_indent_pt is not None:
            paragraph.paragraph_format.left_indent = Pt(float(left_indent_pt))
        elif left_indent_px is not None:
            paragraph.paragraph_format.left_indent = Pt(float(left_indent_px) * 72 / 96)
        if text_indent_pt is not None:
            paragraph.paragraph_format.first_line_indent = Pt(float(text_indent_pt))
        elif text_indent_px is not None:
            paragraph.paragraph_format.first_line_indent = Pt(float(text_indent_px) * 72 / 96)

    @staticmethod
    def _is_caption_paragraph(text: str, format_data: dict[str, Any]) -> bool:
        style_name = str(format_data.get("styleName") or "").strip().upper()
        alignment = format_data.get("alignment")
        if style_name == "TF":
            return True
        if alignment is not None and int(alignment) == int(WD_ALIGN_PARAGRAPH.CENTER):
            return True
        return bool(re.match(r"^(Figure|Table)\s+[A-Za-z0-9.\-]+:\s+\S+", text.strip()))

    @staticmethod
    def _index_notes(notes: list[dict[str, Any]]) -> dict[str, Any]:
        clause_notes: dict[str, list[dict[str, Any]]] = {}
        selection_notes: dict[tuple[str, int], list[dict[str, Any]]] = {}
        for note in notes or []:
            translation = str(note.get("translation") or "").strip()
            if not translation:
                continue
            note_type = str(note.get("type") or "")
            clause_key = str(note.get("clauseKey") or "")
            if note_type == "clause" and clause_key:
                clause_notes.setdefault(clause_key, []).append(note)
            elif note_type == "selection" and clause_key:
                block_index = int(note.get("blockIndex") or 0)
                selection_notes.setdefault((clause_key, block_index), []).append(note)
        return {"clause": clause_notes, "selection": selection_notes}

    @staticmethod
    def _index_highlights(highlights: list[dict[str, Any]]) -> dict[tuple[str, int], list[dict[str, Any]]]:
        indexed: dict[tuple[str, int], list[dict[str, Any]]] = {}
        for item in highlights or []:
            clause_key = str(item.get("clauseKey") or "").strip()
            if not clause_key:
                continue
            block_index = int(item.get("blockIndex", -1))
            if block_index < 0:
                continue
            indexed.setdefault((clause_key, block_index), []).append(item)
        return indexed

    def _add_paragraph_with_notes(self, document: Document, text: str, notes: list[dict[str, Any]]) -> Paragraph:
        paragraph = document.add_paragraph()
        comment_targets = self._build_comment_targets(text, notes)
        if not comment_targets:
            paragraph.add_run(text)
            return paragraph

        cursor = 0
        run_groups: list[tuple[list[Any], dict[str, Any]]] = []
        for start, end, note in comment_targets:
            if start > cursor:
                paragraph.add_run(text[cursor:start])
            commented_run = paragraph.add_run(text[start:end])
            run_groups.append(([commented_run], note))
            cursor = end
        if cursor < len(text):
            paragraph.add_run(text[cursor:])

        for runs, note in run_groups:
            self._add_comment(runs, note)
        return paragraph

    @staticmethod
    def _build_comment_targets(text: str, notes: list[dict[str, Any]]) -> list[tuple[int, int, dict[str, Any]]]:
        ranges: list[tuple[int, int, dict[str, Any]]] = []
        occupied: list[tuple[int, int]] = []
        for note in notes:
            source_text = str(note.get("sourceText") or "").strip()
            if not source_text:
                continue
            start = text.find(source_text)
            if start < 0:
                continue
            end = start + len(source_text)
            if any(not (end <= left or start >= right) for left, right in occupied):
                continue
            ranges.append((start, end, note))
            occupied.append((start, end))
        return sorted(ranges, key=lambda item: item[0])

    def _attach_comment_to_paragraphs(self, paragraphs: list[Paragraph], note: dict[str, Any]) -> None:
        source_text = str(note.get("sourceText") or "").strip()
        if source_text:
            for paragraph in paragraphs:
                if source_text and source_text in paragraph.text:
                    rebuilt = self._rebuild_paragraph_with_single_comment(paragraph, source_text, note)
                    if rebuilt:
                        return
        if paragraphs:
            self._attach_comment_to_paragraph(paragraphs[0], note)

    def _rebuild_paragraph_with_single_comment(self, paragraph: Paragraph, source_text: str, note: dict[str, Any]) -> bool:
        text = paragraph.text
        start = text.find(source_text)
        if start < 0:
            return False
        end = start + len(source_text)
        paragraph.clear()
        runs = []
        if start > 0:
            paragraph.add_run(text[:start])
        comment_run = paragraph.add_run(text[start:end])
        runs.append(comment_run)
        if end < len(text):
            paragraph.add_run(text[end:])
        self._add_comment(runs, note)
        return True

    def _attach_comment_to_paragraph(self, paragraph: Paragraph, note: dict[str, Any]) -> None:
        if not paragraph.runs:
            paragraph.add_run(paragraph.text or " ")
        self._add_comment([paragraph.runs[0]], note)

    @classmethod
    def _compare_clause_nodes(cls, left: dict[str, Any], right: dict[str, Any]) -> int:
        left_path = [str(part) for part in (left.get("clausePath") or [left.get("clauseId") or ""]) if str(part)]
        right_path = [str(part) for part in (right.get("clausePath") or [right.get("clauseId") or ""]) if str(part)]
        path_length = max(len(left_path), len(right_path))
        for index in range(path_length):
            part_compare = cls._compare_clause_part(left_path[index] if index < len(left_path) else "", right_path[index] if index < len(right_path) else "")
            if part_compare != 0:
                return part_compare
        left_order = int(left.get("orderInSource") or 0)
        right_order = int(right.get("orderInSource") or 0)
        if left_order != right_order:
            return left_order - right_order
        return str(left.get("clauseId") or "").locale_compare(str(right.get("clauseId") or "")) if False else (
            -1 if str(left.get("clauseId") or "") < str(right.get("clauseId") or "") else 1 if str(left.get("clauseId") or "") > str(right.get("clauseId") or "") else 0
        )

    @classmethod
    def _compare_clause_part(cls, left: str, right: str) -> int:
        left_tokens = str(left).split(".")
        right_tokens = str(right).split(".")
        length = max(len(left_tokens), len(right_tokens))
        for index in range(length):
            token_compare = cls._compare_mixed_token(left_tokens[index] if index < len(left_tokens) else "", right_tokens[index] if index < len(right_tokens) else "")
            if token_compare != 0:
                return token_compare
        return 0

    @staticmethod
    def _compare_mixed_token(left: str, right: str) -> int:
        left_match = re.match(r"^(\d+)(.*)$", str(left))
        right_match = re.match(r"^(\d+)(.*)$", str(right))
        if left_match and right_match:
            number_compare = int(left_match.group(1)) - int(right_match.group(1))
            if number_compare != 0:
                return number_compare
            left_suffix = left_match.group(2)
            right_suffix = right_match.group(2)
            if left_suffix < right_suffix:
                return -1
            if left_suffix > right_suffix:
                return 1
            return 0
        if str(left) < str(right):
            return -1
        if str(left) > str(right):
            return 1
        return 0

    @staticmethod
    def _add_comment(runs: list[Any], note: dict[str, Any]) -> None:
        translation = str(note.get("translation") or "").strip()
        if not translation or not runs:
            return
        try:
            document = runs[0]._parent.part.document
            comment = document.add_comment(runs, text=translation, author="SpecBot", initials="SB")
            DocxExportService._apply_comment_fonts(comment)
        except Exception:
            return

    @staticmethod
    def _apply_paragraph_highlight(paragraph: Paragraph) -> None:
        if not paragraph.runs:
            paragraph.add_run(paragraph.text or " ")
        for run in paragraph.runs:
            r_pr = run._element.get_or_add_rPr()
            highlight = r_pr.find(qn("w:highlight"))
            if highlight is None:
                highlight = OxmlElement("w:highlight")
                r_pr.append(highlight)
            highlight.set(qn("w:val"), DOCX_HIGHLIGHT_COLOR)

    @staticmethod
    def _apply_cell_highlight(cell: Any) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), DOCX_SHADE_FILL)

    @staticmethod
    def _apply_comment_fonts(comment: Any) -> None:
        preferred_font = "Malgun Gothic"
        for paragraph in getattr(comment, "paragraphs", []):
            for run in getattr(paragraph, "runs", []):
                if not str(getattr(run, "text", "") or "").strip():
                    continue
                try:
                    run.font.name = preferred_font
                    r_pr = run._element.get_or_add_rPr()
                    r_fonts = r_pr.get_or_add_rFonts()
                    r_fonts.set(qn("w:ascii"), preferred_font)
                    r_fonts.set(qn("w:hAnsi"), preferred_font)
                    r_fonts.set(qn("w:eastAsia"), preferred_font)
                    r_fonts.set(qn("w:cs"), preferred_font)
                except Exception:
                    continue


class LLMActionService:
    def __init__(
        self,
        provider: str = "openai",
        model: str = "gpt-4o-mini",
        system_prompt_path: str | Path | None = None,
        user_prompt_path: str | Path | None = None,
        selection_system_prompt_path: str | Path | None = None,
        selection_user_prompt_path: str | Path | None = None,
        max_concurrent_requests: int = 2,
        max_queued_requests: int = 5,
    ) -> None:
        self._provider = provider
        self._model = model
        self._max_concurrent_requests = max(1, int(max_concurrent_requests))
        self._max_queued_requests = max(0, int(max_queued_requests))
        self._queue_condition = threading.Condition()
        self._active_requests = 0
        self._queued_requests = 0
        self._system_prompt = self._load_prompt(
            system_prompt_path,
            default=(
                "You are a professional 3GPP technical translator.\n"
                "Translate the provided context faithfully into Korean.\n"
                "Return only the translated text."
            ),
        )
        self._user_prompt_template = self._load_prompt(
            user_prompt_path,
            default="Translate the following 3GPP context into Korean.\n\n[CONTEXT]\n{context_text}",
        )
        self._selection_system_prompt = self._load_prompt(
            selection_system_prompt_path,
            default=self._system_prompt,
        )
        self._selection_user_prompt_template = self._load_prompt(
            selection_user_prompt_path,
            default=self._user_prompt_template,
        )

    def available_actions(self) -> list[dict[str, str]]:
        return [{"type": "translate", "label": "Translate"}]

    def run(
        self,
        *,
        action_type: str,
        text: str,
        source_language: str,
        target_language: str,
        context: str | None = None,
        action_scope: str | None = None,
        should_cancel=None,
    ) -> dict[str, Any]:
        cleaned_text = text.strip()
        if not cleaned_text:
            raise ValueError("Select text before requesting an LLM action.")
        if action_type != "translate":
            raise ValueError(f"Unsupported action type: {action_type}")

        if self._provider == "openai":
            if not os.environ.get("OPENAI_API_KEY"):
                raise RuntimeError("OPENAI_API_KEY is not configured on the server.")
            return self._run_openai_translation(
                text=cleaned_text,
                source_language=source_language,
                target_language=target_language,
                context=context or "",
                action_scope=action_scope,
                should_cancel=should_cancel,
            )
        if self._provider == "mock":
            if should_cancel and should_cancel():
                raise LLMActionCancelledError("LLM action cancelled by client.")
            return self._run_mock_translation(
                text=cleaned_text,
                source_language=source_language,
                target_language=target_language,
            )
        raise RuntimeError(f"Unsupported LLM provider: {self._provider}")

    def run_limited(
        self,
        *,
        action_type: str,
        text: str,
        source_language: str,
        target_language: str,
        context: str | None = None,
        action_scope: str | None = None,
        should_cancel=None,
    ) -> dict[str, Any]:
        self._acquire_request_slot()
        try:
            return self.run(
                action_type=action_type,
                text=text,
                source_language=source_language,
                target_language=target_language,
                context=context,
                action_scope=action_scope,
                should_cancel=should_cancel,
            )
        finally:
            self._release_request_slot()

    def _acquire_request_slot(self) -> None:
        with self._queue_condition:
            if self._active_requests < self._max_concurrent_requests:
                self._active_requests += 1
                return
            if self._queued_requests >= self._max_queued_requests:
                raise LLMActionQueueFullError(
                    "The translation queue is full. Wait for current requests to finish and try again."
                )
            self._queued_requests += 1
            try:
                while self._active_requests >= self._max_concurrent_requests:
                    self._queue_condition.wait()
                self._queued_requests -= 1
                self._active_requests += 1
            except BaseException:
                self._queued_requests = max(0, self._queued_requests - 1)
                self._queue_condition.notify(1)
                raise

    def _release_request_slot(self) -> None:
        with self._queue_condition:
            self._active_requests = max(0, self._active_requests - 1)
            self._queue_condition.notify(1)

    def _run_openai_translation(
        self,
        *,
        text: str,
        source_language: str,
        target_language: str,
        context: str,
        action_scope: str | None = None,
        should_cancel=None,
    ) -> dict[str, Any]:
        try:
            from langchain_openai import ChatOpenAI
        except Exception as exc:  # pragma: no cover
            raise RuntimeError(f"OpenAI provider is unavailable: {exc}") from exc

        client = ChatOpenAI(model=self._model, temperature=0)
        translated_parts: list[str] = []
        system_prompt, user_prompt_template = self._get_translation_prompt_pair(action_scope)
        for chunk in self._split_translation_text(text):
            if should_cancel and should_cancel():
                raise LLMActionCancelledError("LLM action cancelled by client.")
            context_text = chunk if not context else f"{context}\n\n{chunk}"
            human_prompt = user_prompt_template.replace("{context_text}", context_text)
            response = client.invoke([("system", system_prompt), ("human", human_prompt)])
            content = response.content if isinstance(response.content, str) else str(response.content)
            translated_parts.append(content.strip())
        return {
            "actionType": "translate",
            "provider": "openai",
            "model": self._model,
            "outputText": "\n\n".join(part for part in translated_parts if part),
            "sourceLanguage": source_language,
            "targetLanguage": target_language,
        }

    def _get_translation_prompt_pair(self, action_scope: str | None) -> tuple[str, str]:
        if str(action_scope or "").strip().lower() == "selection":
            return self._selection_system_prompt, self._selection_user_prompt_template
        return self._system_prompt, self._user_prompt_template

    @staticmethod
    def _load_prompt(path: str | Path | None, *, default: str) -> str:
        if not path:
            return default
        prompt_path = Path(path)
        try:
            return prompt_path.read_text(encoding="utf-8").strip() or default
        except OSError:
            return default

    @staticmethod
    def _split_translation_text(text: str, limit: int = TRANSLATION_CHUNK_LIMIT) -> list[str]:
        cleaned = text.strip()
        if len(cleaned) <= limit:
            return [cleaned]

        paragraphs = [part.strip() for part in re.split(r"\n{2,}", cleaned) if part.strip()]
        if not paragraphs:
            paragraphs = [cleaned]

        chunks: list[str] = []
        current_parts: list[str] = []
        current_length = 0

        for paragraph in paragraphs:
            if len(paragraph) > limit:
                if current_parts:
                    chunks.append("\n\n".join(current_parts))
                    current_parts = []
                    current_length = 0
                chunks.append(paragraph)
                continue

            separator = 2 if current_parts else 0
            next_length = current_length + separator + len(paragraph)
            if current_parts and next_length > limit:
                chunks.append("\n\n".join(current_parts))
                current_parts = [paragraph]
                current_length = len(paragraph)
            else:
                current_parts.append(paragraph)
                current_length = next_length

        if current_parts:
            chunks.append("\n\n".join(current_parts))

        return [chunk for chunk in chunks if chunk]

    @staticmethod
    def _run_mock_translation(*, text: str, source_language: str, target_language: str) -> dict[str, Any]:
        return {
            "actionType": "translate",
            "provider": "mock",
            "model": "mock-translation",
            "outputText": f"[{source_language}->{target_language}] {text}",
            "sourceLanguage": source_language,
            "targetLanguage": target_language,
        }


def sanitize_file_stem(value: str) -> str:
    candidate = INVALID_FILENAME_CHARS.sub("", value.strip())
    candidate = WHITESPACE_RUN.sub("_", candidate)
    candidate = candidate.strip(".- ")
    return candidate or "clause-export"


@dataclass(frozen=True)
class SpecbotQueryDefaults:
    base_url: str = "http://localhost:8080"
    config_base_url: str = "http://localhost:19071"
    limit: int = 4
    iterations: int = 1
    next_iteration_limit: int = 2
    followup_mode: str = "sentence-summary"
    summary: str = "short"
    registry: str = "./artifacts/spec_query_registry.json"
    local_model_dir: str = "models/Qwen3-Embedding-0.6B"
    device: str = "cuda"
    sparse_boost: float = 0.0
    vector_boost: float = 1.0

    def to_dict(self) -> dict[str, object]:
        return {
            "baseUrl": self.base_url,
            "configBaseUrl": self.config_base_url,
            "limit": self.limit,
            "iterations": self.iterations,
            "nextIterationLimit": self.next_iteration_limit,
            "followupMode": self.followup_mode,
            "summary": self.summary,
            "registry": self.registry,
            "localModelDir": self.local_model_dir,
            "device": self.device,
            "sparseBoost": self.sparse_boost,
            "vectorBoost": self.vector_boost,
        }


class SpecbotQueryService:
    def __init__(self, project_root: str | Path, defaults: SpecbotQueryDefaults | None = None) -> None:
        self._project_root = Path(project_root)
        self._defaults = defaults or SpecbotQueryDefaults()

    @property
    def defaults(self) -> SpecbotQueryDefaults:
        return self._defaults

    def run(
        self,
        query: str,
        settings: dict[str, Any] | None = None,
        exclude_specs: list[str] | None = None,
        exclude_clauses: list[dict[str, Any]] | None = None,
    ) -> dict[str, Any]:
        cleaned_query = query.strip()
        if not cleaned_query:
            raise ValueError("Query is required.")

        effective = self._merge_settings(settings or {})
        command = self._build_command(cleaned_query, effective)
        result = subprocess.run(
            command,
            cwd=self._project_root,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            check=False,
            timeout=180,
        )
        if result.returncode != 0:
            detail = (result.stderr or result.stdout).strip() or f"SpecBot exited with code {result.returncode}"
            raise RuntimeError(detail)

        payload = self._parse_json_output(result.stdout)
        return {
            "query": cleaned_query,
            "settings": effective,
            "hits": self._extract_hits(payload),
            "rawResult": payload,
            "command": self._display_command(cleaned_query, effective),
        }

    def _merge_settings(self, settings: dict[str, Any]) -> dict[str, Any]:
        defaults = self._defaults.to_dict()
        merged = {**defaults, **{key: value for key, value in settings.items() if value is not None}}
        merged["limit"] = max(1, int(merged["limit"]))
        merged["iterations"] = max(1, int(merged["iterations"]))
        merged["nextIterationLimit"] = max(1, int(merged["nextIterationLimit"]))
        merged["followupMode"] = str(merged.get("followupMode") or self._defaults.followup_mode).strip() or self._defaults.followup_mode
        merged["sparseBoost"] = float(merged["sparseBoost"])
        merged["vectorBoost"] = float(merged["vectorBoost"])
        return merged

    def _build_command(self, query: str, settings: dict[str, Any]) -> list[str]:
        command = [
            sys.executable,
            "-m",
            "app.main",
            "iterative-query-vespa-http",
            "--query",
            query,
            "--base-url",
            str(settings["baseUrl"]),
            "--limit",
            str(settings["limit"]),
            "--iterations",
            str(settings["iterations"]),
            "--next-iteration-limit",
            str(settings["nextIterationLimit"]),
            "--followup-mode",
            str(settings["followupMode"]),
            "--summary",
            str(settings["summary"]),
            "--registry",
            str(settings["registry"]),
            "--local-model-dir",
            str(settings["localModelDir"]),
            "--device",
            str(settings["device"]),
            "--sparse-boost",
            str(settings["sparseBoost"]),
            "--vector-boost",
            str(settings["vectorBoost"]),
        ]
        if str(settings.get("configBaseUrl", "")).strip():
            command.extend(["--config-base-url", str(settings["configBaseUrl"])])
        return command

    @staticmethod
    def _parse_json_output(stdout: str) -> dict[str, Any]:
        text = stdout.strip()
        if not text:
            raise RuntimeError("SpecBot returned no output.")
        try:
            return dict(json.loads(text))
        except Exception as exc:
            raise RuntimeError(f"SpecBot returned invalid JSON: {exc}") from exc

    @staticmethod
    def _extract_hits(payload: dict[str, Any]) -> list[dict[str, Any]]:
        hits: list[dict[str, Any]] = []
        for item in payload.get("relevant_documents", []) or []:
            texts = item.get("texts") or []
            preview = texts[0][:240] if texts else ""
            hits.append(
                {
                    "specNo": str(item.get("spec_no", "")),
                    "clauseId": str(item.get("clause_id", "")),
                    "parentClauseId": str(item.get("parent_clause_id", "")),
                    "clausePath": [str(part) for part in item.get("clause_path", [])],
                    "textPreview": preview,
                }
            )
        return hits

    def _display_command(self, query: str, settings: dict[str, Any]) -> str:
        return " ".join(
            [
                "python3 -m app.main iterative-query-vespa-http",
                f'--query "{query}"',
                f'--base-url {settings["baseUrl"]}',
                f'--config-base-url {settings["configBaseUrl"]}',
                f'--limit {settings["limit"]}',
                f'--iterations {settings["iterations"]}',
                f'--next-iteration-limit {settings["nextIterationLimit"]}',
                f'--followup-mode {settings["followupMode"]}',
                f'--summary {settings["summary"]}',
                f'--registry {settings["registry"]}',
                f'--local-model-dir {settings["localModelDir"]}',
                f'--device {settings["device"]}',
                f'--sparse-boost {settings["sparseBoost"]}',
                f'--vector-boost {settings["vectorBoost"]}',
            ]
        )


class SpecbotQueryHttpService:
    def __init__(self, base_url: str, defaults: SpecbotQueryDefaults | None = None, timeout_seconds: float = 180.0) -> None:
        self._base_url = base_url.rstrip("/")
        self._defaults = defaults or SpecbotQueryDefaults()
        self._timeout_seconds = timeout_seconds

    @property
    def defaults(self) -> SpecbotQueryDefaults:
        return self._defaults

    def run(
        self,
        query: str,
        settings: dict[str, Any] | None = None,
        exclude_specs: list[str] | None = None,
        exclude_clauses: list[dict[str, Any]] | None = None,
        release_data: str | None = None,
        release: str | None = None,
    ) -> dict[str, Any]:
        payload = json.dumps(
            {
                "query": query,
                "releaseData": release_data,
                "release": release,
                "settings": settings or self._defaults.to_dict(),
                "excludeSpecs": exclude_specs or [],
                "excludeClauses": exclude_clauses or [],
            },
            ensure_ascii=True,
        ).encode("utf-8")
        request = urllib_request.Request(
            url=f"{self._base_url}/query",
            data=payload,
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib_request.urlopen(request, timeout=self._timeout_seconds) as response:
                body = response.read().decode("utf-8")
        except urllib_error.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="replace")
            raise RuntimeError(_extract_error_detail(detail) or f"Query API HTTP {exc.code}") from exc
        except urllib_error.URLError as exc:
            raise RuntimeError(f"Unable to reach SpecBot query API at {self._base_url}: {exc.reason}") from exc

        try:
            return dict(json.loads(body))
        except Exception as exc:
            raise RuntimeError(f"SpecBot query API returned invalid JSON: {exc}") from exc

    async def run_async(
        self,
        query: str,
        settings: dict[str, Any] | None = None,
        exclude_specs: list[str] | None = None,
        exclude_clauses: list[dict[str, Any]] | None = None,
        release_data: str | None = None,
        release: str | None = None,
        should_cancel=None,
    ) -> dict[str, Any]:
        import httpx

        payload = {
            "query": query,
            "releaseData": release_data,
            "release": release,
            "settings": settings or self._defaults.to_dict(),
            "excludeSpecs": exclude_specs or [],
            "excludeClauses": exclude_clauses or [],
        }

        async with httpx.AsyncClient(timeout=self._timeout_seconds) as client:
            request_task = asyncio.create_task(client.post(f"{self._base_url}/query", json=payload))
            try:
                while True:
                    if should_cancel and should_cancel():
                        request_task.cancel()
                        raise RuntimeError("SpecBot query cancelled by client.")
                    if request_task.done():
                        break
                    await asyncio.sleep(0.2)
                response = await request_task
            except asyncio.CancelledError as exc:
                request_task.cancel()
                raise RuntimeError("SpecBot query cancelled by client.") from exc
            except httpx.HTTPError as exc:
                raise RuntimeError(f"Unable to reach SpecBot query API at {self._base_url}: {exc}") from exc

        if response.status_code >= 400:
            if response.status_code == 429:
                raise LLMActionQueueFullError(
                    _extract_error_detail(response.text) or "The shared query/translation queue is full."
                )
            if response.status_code == 499:
                raise LLMActionCancelledError(_extract_error_detail(response.text) or "SpecBot query cancelled by client.")
            raise RuntimeError(_extract_error_detail(response.text) or f"Query API HTTP {response.status_code}")

        try:
            return dict(response.json())
        except Exception as exc:
            raise RuntimeError(f"SpecBot query API returned invalid JSON: {exc}") from exc


class LLMActionHttpService:
    def __init__(self, base_url: str, timeout_seconds: float = 180.0) -> None:
        self._base_url = base_url.rstrip("/")
        self._timeout_seconds = timeout_seconds

    def available_actions(self) -> list[dict[str, str]]:
        return [{"type": "translate", "label": "Translate"}]

    async def run_async(
        self,
        *,
        action_type: str,
        action_scope: str | None = None,
        text: str,
        source_language: str,
        target_language: str,
        context: str | None = None,
        should_cancel=None,
    ) -> dict[str, Any]:
        import httpx

        payload = {
            "actionType": action_type,
            "actionScope": action_scope,
            "text": text,
            "sourceLanguage": source_language,
            "targetLanguage": target_language,
            "context": context,
        }

        async with httpx.AsyncClient(timeout=self._timeout_seconds) as client:
            request_task = asyncio.create_task(client.post(f"{self._base_url}/llm-actions", json=payload))
            try:
                while True:
                    if should_cancel and should_cancel():
                        request_task.cancel()
                        raise LLMActionCancelledError("LLM action cancelled by client.")
                    if request_task.done():
                        break
                    await asyncio.sleep(0.2)
                response = await request_task
            except asyncio.CancelledError as exc:
                request_task.cancel()
                raise LLMActionCancelledError("LLM action cancelled by client.") from exc
            except httpx.HTTPError as exc:
                raise RuntimeError(f"Unable to reach SpecBot query API at {self._base_url}: {exc}") from exc

        if response.status_code == 429:
            raise LLMActionQueueFullError(_extract_error_detail(response.text) or "LLM action queue is full.")
        if response.status_code == 499:
            raise LLMActionCancelledError(_extract_error_detail(response.text) or "LLM action cancelled by client.")
        if response.status_code >= 400:
            raise RuntimeError(_extract_error_detail(response.text) or f"LLM action API HTTP {response.status_code}")

        try:
            return dict(response.json())
        except Exception as exc:
            raise RuntimeError(f"LLM action API returned invalid JSON: {exc}") from exc


def _extract_error_detail(payload_text: str) -> str:
    text = (payload_text or "").strip()
    if not text:
        return ""
    try:
        payload = json.loads(text)
    except Exception:
        return text
    return _stringify_error_detail(payload)


def _stringify_error_detail(value: Any) -> str:
    if isinstance(value, dict):
        if "detail" in value:
            return _stringify_error_detail(value["detail"])
        if "message" in value:
            return _stringify_error_detail(value["message"])
        parts = [f"{key}: {_stringify_error_detail(item)}" for key, item in value.items()]
        return "; ".join(part for part in parts if part)
    if isinstance(value, list):
        parts = [_stringify_error_detail(item) for item in value]
        return "; ".join(part for part in parts if part)
    if value is None:
        return ""
    return str(value)
